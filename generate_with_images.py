#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo file Word Chương 2 và Chương 3 với HÌnh ảnh thực tế (PNG/JPG)
"""

import os, glob
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths to generated images ───
BRAIN_DIR = '/Users/nguyennghia/.gemini/antigravity/brain/cfb9d904-e623-4ba7-9b16-5b1cbcf4b809'

def find_image(prefix):
    """Find latest generated image by prefix."""
    pattern = os.path.join(BRAIN_DIR, f'{prefix}_*.jpg')
    files = glob.glob(pattern)
    if files:
        return sorted(files)[-1]
    # Try png
    pattern = os.path.join(BRAIN_DIR, f'{prefix}_*.png')
    files = glob.glob(pattern)
    return sorted(files)[-1] if files else None

IMAGES = {
    'usecase': find_image('usecase_diagram'),
    'fdd': find_image('fdd_diagram'),
    'architecture': find_image('system_architecture'),
    'seq_login': find_image('seq_login'),
    'seq_checkout': find_image('seq_checkout'),
    'seq_review': find_image('seq_review'),
    'seq_admin_product': find_image('seq_admin_product'),
    'seq_livechat': find_image('seq_livechat'),
    'seq_lucky_wheel': find_image('seq_lucky_wheel'),
    'erd_main': find_image('erd_main'),
    'erd_extended': find_image('erd_extended'),
    'dfd': find_image('dfd_context'),
}

print("Images found:")
for k, v in IMAGES.items():
    print(f"  {k}: {v}")

# ─────────────────────────── Helpers ───────────────────────────

def set_font(run, bold=False, size=12, color=None, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)


def add_heading(doc, text, level=1, size=14, bold=True, color=None, space_before=12, space_after=6):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    set_font(run, bold=bold, size=size, color=color)
    return para


def add_para(doc, text, indent=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, space_before=2, space_after=4):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if indent:
        para.paragraph_format.left_indent = Pt(indent)
    run = para.add_run(text)
    set_font(run, bold=bold, size=size)
    return para


def add_bullet(doc, text, level=0, size=12):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Pt(18 + level * 18)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    set_font(run, size=size)
    return para


def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '6')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_table_header_row(table, headers, bg_color='1F3864'):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        set_font(run, bold=True, size=11, color=(255, 255, 255))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color)
        tcPr.append(shd)


def add_figure_caption(doc, text, size=11):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(10)
    run = para.add_run(text)
    set_font(run, bold=True, size=size, italic=True, color=(80, 80, 80))
    return para


def add_image(doc, img_key, width_cm=15):
    """Insert an actual image into doc."""
    img_path = IMAGES.get(img_key)
    if img_path and os.path.exists(img_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        return True
    else:
        print(f"  ⚠️  Image not found: {img_key}")
        return False


def add_steps(doc, steps_data, title=None):
    """Add numbered step-by-step description of a diagram."""
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        set_font(r, bold=True, size=11, italic=True, color=(31, 56, 100))
    for step_num, step_text in enumerate(steps_data, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.left_indent = Pt(18)
        # Step number
        rn = para.add_run(f'Bước {step_num}: ')
        set_font(rn, bold=True, size=11, color=(31, 56, 100))
        # Step content
        rc = para.add_run(step_text)
        set_font(rc, size=11)


def add_page_break(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
#  CHAPTER 2
# ═══════════════════════════════════════════════════════════════

def build_chapter2(doc):
    add_heading(doc, 'CHƯƠNG 2: ĐẶC TẢ YÊU CẦU', size=16, color=(31, 56, 100), space_before=18, space_after=12)

    # 2.1
    add_heading(doc, '2.1. Khảo sát hiện trạng và bài toán đặt ra', size=14, color=(31, 56, 100))
    add_para(doc, 'Trong bối cảnh thương mại điện tử phát triển mạnh mẽ, nhu cầu mua sắm laptop trực tuyến ngày càng tăng cao. Nhiều website bán laptop hiện nay vẫn còn hạn chế về tính năng quản lý, trải nghiệm người dùng và bảo mật. Đề tài hướng đến xây dựng một hệ thống bán laptop trực tuyến hoàn chỉnh với đầy đủ chức năng hiện đại, đáp ứng nhu cầu thực tế của người mua lẫn quản trị viên.')
    add_para(doc, 'Hệ thống được xây dựng gồm hai phần chính: giao diện người dùng (Frontend sử dụng Next.js/React) và hệ thống quản trị (Admin Panel), kết hợp với backend Laravel xử lý toàn bộ nghiệp vụ và API RESTful. Cơ sở dữ liệu MySQL lưu trữ toàn bộ thông tin sản phẩm, đơn hàng, người dùng, đánh giá và các tính năng mở rộng như vòng quay may mắn, live chat, coupon giảm giá và tích hợp AI Chatbot.')

    # 2.2
    add_heading(doc, '2.2. Phân tích yêu cầu chức năng', size=14, color=(31, 56, 100))
    add_para(doc, 'Hệ thống được phân tích thành các nhóm chức năng chính theo từng đối tượng sử dụng:')

    add_heading(doc, '2.2.1. Người dùng chưa đăng nhập (Guest)', size=13, color=(70, 100, 160))
    for item in [
        'Xem danh sách sản phẩm: duyệt tất cả sản phẩm, lọc theo danh mục, sắp xếp theo giá/đánh giá.',
        'Tìm kiếm sản phẩm: tìm theo từ khóa tên, mô tả sản phẩm.',
        'Xem chi tiết sản phẩm: hình ảnh lớn (đa ảnh), mô tả đầy đủ, tồn kho, đánh giá từ người mua.',
        'Thêm vào danh sách yêu thích: lưu vào localStorage không cần đăng nhập.',
        'Đăng ký tài khoản: nhập thông tin cá nhân để tạo tài khoản mới.',
        'Đăng nhập: xác thực bằng email/mật khẩu, nhận Sanctum token.',
        'Gửi liên hệ/hỏi đáp: qua form contact hoặc live chat với admin.',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '2.2.2. Người dùng đã đăng nhập (User)', size=13, color=(70, 100, 160))
    for item in [
        'Tất cả chức năng của Guest.',
        'Quản lý giỏ hàng: thêm, cập nhật số lượng, xóa sản phẩm khỏi giỏ hàng.',
        'Đặt hàng: chọn địa chỉ nhận hàng, phương thức thanh toán (COD, VNPay, MoMo), nhập mã giảm giá (Coupon).',
        'Vòng quay may mắn: quay 1 lần/ngày để nhận coupon giảm giá ngẫu nhiên.',
        'Xem lịch sử đơn hàng: danh sách đơn, trạng thái từng đơn (pending, processing, shipping, completed, cancelled).',
        'Đánh giá sản phẩm: chỉ cho phép đánh giá sau khi đơn hàng đã giao thành công (status = completed).',
        'Quản lý tài khoản cá nhân: cập nhật tên, số điện thoại, địa chỉ; đổi mật khẩu.',
        'Live Chat hỗ trợ: gửi tin nhắn trực tiếp với admin/tư vấn viên theo thời gian thực.',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '2.2.3. Quản trị viên (Admin)', size=13, color=(70, 100, 160))
    for item in [
        'Đăng nhập Admin: truy cập trang quản trị riêng biệt với phân quyền admin.',
        'Quản lý sản phẩm: thêm, sửa, xóa sản phẩm; upload nhiều ảnh; quản lý tồn kho, danh mục.',
        'Quản lý danh mục: thêm, sửa, xóa danh mục sản phẩm.',
        'Quản lý đơn hàng: xem danh sách, chi tiết từng đơn, cập nhật trạng thái, hủy đơn hàng.',
        'Quản lý người dùng: xem danh sách, phân quyền user/admin, xóa tài khoản.',
        'Thống kê & báo cáo: doanh thu theo ngày/tháng, top sản phẩm bán chạy, số đơn hàng theo trạng thái, xuất PDF.',
        'Quản lý coupon: tạo, sửa, xóa mã giảm giá (loại phần trăm hoặc cố định); giới hạn lượt dùng, thời hạn sử dụng.',
        'Quản lý live chat: xem các session chat, trả lời khách hàng, đóng session.',
        'Quản lý liên hệ/hỗ trợ: xem danh sách yêu cầu liên hệ, xử lý, ghi chú trạng thái.',
        'Cài đặt trang chủ: cập nhật banner, slider, nội dung trang chủ.',
        'Cài đặt thanh toán: cấu hình VNPay, MoMo.',
    ]:
        add_bullet(doc, item)

    # 2.3 Yêu cầu phi chức năng
    add_heading(doc, '2.3. Yêu cầu phi chức năng', size=14, color=(31, 56, 100))
    nfr = [
        ('Hiệu năng', 'API phản hồi dưới 500ms; frontend tải trang dưới 3 giây.'),
        ('Bảo mật', 'Xác thực qua Laravel Sanctum; phân quyền admin/user; bcrypt; CSRF protection.'),
        ('Khả năng mở rộng', 'RESTful API dễ tích hợp mobile app, đối tác thanh toán, vận chuyển.'),
        ('Giao diện & UX', 'Responsive mọi thiết bị; toast/popup thông báo; phân trang, lọc, sắp xếp.'),
        ('Toàn vẹn dữ liệu', 'Ràng buộc FK trong CSDL; validate server-side; transaction rollback.'),
        ('Khả năng bảo trì', 'Kiến trúc MVC rõ ràng; comment đầy đủ; dễ nâng cấp tính năng.'),
    ]
    table = doc.add_table(rows=len(nfr)+1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    add_table_header_row(table, ['Yêu cầu phi chức năng', 'Mô tả chi tiết'])
    for i, (req, desc) in enumerate(nfr):
        row = table.rows[i+1]
        row.cells[0].text = ''
        row.cells[1].text = ''
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(req)
        set_font(r0, bold=True, size=11)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(desc)
        set_font(r1, size=11)
    set_table_borders(table)
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(11.5)
    add_figure_caption(doc, 'Bảng 2.1 – Yêu cầu phi chức năng của hệ thống')

    # 2.4 Use Case
    add_heading(doc, '2.4. Mô hình Use Case (Tổng quát)', size=14, color=(31, 56, 100))
    add_para(doc, 'Sơ đồ Use Case (Use Case Diagram) là công cụ mô hình hóa UML dùng để mô tả các chức năng mà hệ thống cung cấp cho người dùng và cách các tác nhân tương tác với hệ thống. Trong sơ đồ, mỗi hình ellipse (oval) đại diện cho một Use Case (chức năng), hình người (Stick Figure) đại diện cho tác nhân (Actor), đường nối thể hiện quan hệ tham gia giữa tác nhân và chức năng.')
    add_para(doc, 'Hệ thống bán laptop trực tuyến gồm 2 tác nhân chính:')
    add_bullet(doc, 'Người Dùng (User): tác nhân chính phía khách hàng, bao gồm cả Guest (chưa đăng nhập) và User đã đăng nhập.')
    add_bullet(doc, 'Quản Trị Viên (Admin): tác nhân quản lý hệ thống, có toàn quyền trên dữ liệu và cấu hình.')
    add_image(doc, 'usecase', width_cm=15)
    add_figure_caption(doc, 'Hình 2.1 – Sơ đồ Use Case tổng quát của hệ thống')
    add_steps(doc,
        [
            'Người Dùng (User) là tác nhân phía trái: có 12 Use Case từ xem sản phẩm, đặt hàng, đánh giá đến live chat và vòng quay may mắn.',
            'Quản Trị Viên (Admin) là tác nhân phía phải: có 9 Use Case đặc thù quản lý như quản lý sản phẩm, quản lý đơn hàng, xuất báo cáo PDF.',
            'Các Use Case của Người Dùng bao gồm: Xem danh sách sản phẩm → Tìm kiếm & lọc → Xem chi tiết → Quản lý giỏ hàng → Áp dụng coupon → Đặt hàng (COD/VNPay/MoMo).',
            'Các Use Case bổ sung của Người Dùng: Đăng ký / Đăng nhập → Xem lịch sử đơn hàng → Đánh giá sản phẩm → Vòng quay may mắn → Live Chat hỗ trợ → Cập nhật tài khoản cá nhân.',
            'Các Use Case của Admin bao gồm: Quản lý sản phẩm & danh mục → Quản lý đơn hàng → Quản lý người dùng → Thống kê & báo cáo doanh thu.',
            'Các Use Case mở rộng của Admin: Quản lý coupon giảm giá → Quản lý live chat → Quản lý liên hệ khách hàng → Cài đặt trang chủ & thanh toán → Xuất báo cáo PDF.',
            'Hai tác nhân hoàn toàn độc lập: Frontend người dùng và Admin Panel là hai ứng dụng Next.js riêng biệt, đều giao tiếp với Backend Laravel qua RESTful API được bảo mật bằng Laravel Sanctum.',
        ],
        title='Mô tả từng thành phần trong sơ đồ Use Case:'
    )

    # 2.5 FDD
    add_heading(doc, '2.5. Sơ đồ phân rã chức năng (Functional Decomposition Diagram)', size=14, color=(31, 56, 100))
    add_para(doc, 'Sơ đồ phân rã chức năng (FDD – Functional Decomposition Diagram) là kỹ thuật phân tích hệ thống từ trên xuống (Top-Down), mô tả cấu trúc phân cấp các chức năng từ mức tổng quát đến chi tiết. Mỗi nút trong sơ đồ đại diện cho một chức năng hoặc nhóm chức năng, các đường kết nối thể hiện quan hệ bao gồm (Parent – Child). Phương pháp này giúp đảm bảo không bỏ sót chức năng nào trong quá trình phân tích và thiết kế.')
    add_image(doc, 'fdd', width_cm=16)
    add_figure_caption(doc, 'Hình 2.2 – Sơ đồ phân rã chức năng (FDD) của hệ thống')
    add_steps(doc,
        [
            'Nút gốc (Root): "Hệ Thống Bán Laptop Trực Tuyến" — đây là chức năng tổng quát nhất, là điểm xuất phát của toàn bộ sơ đồ phân rã.',
            'Cấp 1 – 8 nhóm chức năng chính: (1) Quản Lý Người Dùng, (2) Quản Lý Sản Phẩm, (3) Quản Lý Giỏ Hàng & Đặt Hàng, (4) Quản Lý Đơn Hàng (Admin), (5) Đánh Giá Sản Phẩm, (6) Khuyến Mãi & Gamification, (7) Hỗ Trợ Khách Hàng, (8) Thống Kê & Báo Cáo.',
            'Nhóm 1 – Quản Lý Người Dùng gồm 4 chức năng con: Đăng ký, Đăng nhập, Cập nhật thông tin cá nhân, Phân quyền (user/admin).',
            'Nhóm 2 – Quản Lý Sản Phẩm gồm 5 chức năng con: Xem/Tìm kiếm, Xem chi tiết, Thêm/Sửa/Xóa (Admin), Upload ảnh, Quản lý danh mục.',
            'Nhóm 3 – Quản Lý Giỏ Hàng & Đặt Hàng gồm 4 chức năng: Thêm/cập nhật giỏ, Áp dụng Coupon, Đặt hàng (COD/VNPay/MoMo), Xem lịch sử.',
            'Nhóm 4 – Quản Lý Đơn Hàng (Admin) gồm 3 chức năng: Xem danh sách, Cập nhật trạng thái, Hủy đơn hàng.',
            'Nhóm 5 – Đánh Giá Sản Phẩm gồm 2 chức năng: Đánh giá sau khi nhận hàng, Xem đánh giá (chỉ cho phép sau khi đơn hàng được giao).',
            'Nhóm 6 – Khuyến Mãi & Gamification gồm 3 chức năng: Tạo/quản lý Coupon, Áp dụng Coupon khi đặt hàng, Vòng quay may mắn (mỗi ngày 1 lần).',
            'Nhóm 7 – Hỗ Trợ Khách Hàng gồm 3 chức năng: Form liên hệ, Live Chat với admin, AI Chatbot tư vấn tự động (Google Gemini).',
            'Nhóm 8 – Thống Kê & Báo Cáo gồm 4 chức năng (chỉ Admin): Doanh thu theo ngày/tháng, Top sản phẩm bán chạy, Số đơn theo trạng thái, Xuất PDF báo cáo.',
        ],
        title='Mô tả từng nhóm chức năng trong sơ đồ FDD:'
    )

    # 2.6 Đặc tả UC
    add_heading(doc, '2.6. Đặc tả Use Case chi tiết', size=14, color=(31, 56, 100))
    add_para(doc, 'Dưới đây là bảng đặc tả chi tiết các Use Case quan trọng nhất của hệ thống:')

    uc_specs = [
        {
            'title': 'UC-01: Đặt hàng',
            'data': [
                ('Mã Use Case', 'UC-01'),
                ('Tên Use Case', 'Đặt hàng'),
                ('Tác nhân', 'Người dùng đã đăng nhập (User)'),
                ('Mô tả', 'Người dùng đặt hàng các sản phẩm trong giỏ hàng'),
                ('Tiền điều kiện', 'User đã đăng nhập; giỏ hàng có ít nhất 1 sản phẩm'),
                ('Hậu điều kiện', 'Đơn hàng được tạo; tồn kho cập nhật; giỏ hàng xóa'),
                ('Dòng sự kiện chính', '1. User vào giỏ hàng\n2. Nhập/chọn địa chỉ nhận hàng\n3. Chọn phương thức thanh toán (COD/VNPay/MoMo)\n4. Áp dụng mã coupon (nếu có)\n5. Nhấn "Đặt hàng" → Hệ thống kiểm tra tồn kho\n6. Tạo đơn hàng, trừ tồn kho, xóa giỏ hàng\n7. Hiển thị xác nhận thành công'),
                ('Dòng sự kiện thay thế', 'Hết hàng → báo lỗi, không tạo đơn\nVNPay/MoMo thất bại → hủy đơn, hoàn kho'),
            ]
        },
        {
            'title': 'UC-02: Đánh giá sản phẩm',
            'data': [
                ('Mã Use Case', 'UC-02'),
                ('Tên Use Case', 'Đánh giá sản phẩm'),
                ('Tác nhân', 'Người dùng đã đăng nhập (User)'),
                ('Mô tả', 'Người dùng đánh giá sản phẩm đã mua sau khi nhận hàng thành công'),
                ('Tiền điều kiện', 'User đã đăng nhập; có đơn hàng trạng thái "completed"'),
                ('Hậu điều kiện', 'Đánh giá lưu vào CSDL; avg_rating sản phẩm cập nhật'),
                ('Dòng sự kiện chính', '1. User vào lịch sử đơn hàng\n2. Chọn đơn đã giao → nhấn "Đánh giá"\n3. Hiển thị form (số sao 1–5, nhận xét)\n4. User nhập và gửi đánh giá\n5. Hệ thống kiểm tra → lưu đánh giá\n6. Thông báo thành công'),
                ('Dòng sự kiện thay thế', 'Đơn chưa giao → từ chối, hiển thị lỗi 403\nĐã đánh giá → hiển thị đánh giá cũ'),
            ]
        },
        {
            'title': 'UC-03: Quản lý đơn hàng (Admin)',
            'data': [
                ('Mã Use Case', 'UC-03'),
                ('Tên Use Case', 'Quản lý đơn hàng'),
                ('Tác nhân', 'Quản trị viên (Admin)'),
                ('Mô tả', 'Admin xem, cập nhật trạng thái và hủy đơn hàng'),
                ('Tiền điều kiện', 'Admin đã đăng nhập với quyền admin'),
                ('Hậu điều kiện', 'Trạng thái đơn hàng được cập nhật trong CSDL'),
                ('Dòng sự kiện chính', '1. Admin truy cập Quản lý Đơn hàng\n2. Hiển thị danh sách + bộ lọc trạng thái\n3. Admin chọn đơn → cập nhật trạng thái mới\n4. Lưu thay đổi → hiển thị cập nhật'),
                ('Dòng sự kiện thay thế', 'Đơn "completed" hoặc "cancelled" → không cho sửa\nAdmin có thể hủy đơn pending/processing'),
            ]
        },
    ]

    for uc in uc_specs:
        add_heading(doc, uc['title'], size=12, color=(70, 100, 160), space_before=10, space_after=4)
        table = doc.add_table(rows=len(uc['data']), cols=2)
        table.style = 'Table Grid'
        set_table_borders(table)
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(12)
        for i, (field, value) in enumerate(uc['data']):
            row = table.rows[i]
            row.cells[0].text = ''
            row.cells[1].text = ''
            tcPr = row.cells[0]._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'DCE6F1')
            tcPr.append(shd)
            p0 = row.cells[0].paragraphs[0]
            r0 = p0.add_run(field)
            set_font(r0, bold=True, size=10)
            p1 = row.cells[1].paragraphs[0]
            r1 = p1.add_run(value)
            set_font(r1, size=10)
        doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
#  CHAPTER 3
# ═══════════════════════════════════════════════════════════════

def build_chapter3(doc):
    add_page_break(doc)
    add_heading(doc, 'CHƯƠNG 3: THIẾT KẾ VÀ CÀI ĐẶT GIẢI PHÁP', size=16, color=(31, 56, 100), space_before=18, space_after=12)

    # 3.1 Kiến trúc
    add_heading(doc, '3.1. Thiết kế kiến trúc tổng thể (System Architecture)', size=14, color=(31, 56, 100))
    add_para(doc, 'Kiến trúc phần mềm (Software Architecture) xác định cấu trúc tổng thể của hệ thống, bao gồm các thành phần chính và cách chúng tương tác với nhau. Hệ thống bán laptop trực tuyến được thiết kế theo mô hình kiến trúc 3 tầng (Three-Tier Architecture) — một trong những mô hình phổ biến và bền vững nhất trong phát triển ứng dụng web hiện đại:')
    add_bullet(doc, 'Tầng Trình Bày (Presentation Layer): Gồm hai ứng dụng Next.js/React độc lập — Frontend người dùng và Admin Panel. Tầng này chịu trách nhiệm hiển thị giao diện, nhận input từ người dùng và gọi API để lấy/gửi dữ liệu. Hai ứng dụng này chạy độc lập, cho phép phân quyền và triển khai riêng biệt.')
    add_bullet(doc, 'Tầng Nghiệp Vụ (Business Logic Layer): Backend Laravel 10 (PHP 8.x) xử lý toàn bộ logic nghiệp vụ: xác thực người dùng (Sanctum), quản lý đơn hàng, tích hợp thanh toán (VNPay, MoMo), AI Chatbot (Gemini API), live chat, quản lý coupon và vòng quay may mắn. Tầng này tuân thủ kiến trúc MVC, phân tách rõ Controller – Model – Route.')
    add_bullet(doc, 'Tầng Dữ Liệu (Data Layer): MySQL Database lưu trữ toàn bộ 15 bảng dữ liệu. Laravel Eloquent ORM đóng vai trò trung gian, cung cấp giao diện hướng đối tượng để thao tác CSDL, đảm bảo toàn vẹn dữ liệu qua ràng buộc khóa ngoại.')
    add_image(doc, 'architecture', width_cm=15.5)
    add_figure_caption(doc, 'Hình 3.1 – Kiến trúc tổng thể hệ thống bán laptop trực tuyến')
    add_para(doc, 'Nhận xét kiến trúc: Việc tách biệt Frontend người dùng và Admin Panel thành hai ứng dụng độc lập giúp tăng tính bảo mật (Admin Panel có thể triển khai trên subdomain riêng) và dễ bảo trì. Backend Laravel đóng vai trò API Gateway duy nhất, cho phép mở rộng sang ứng dụng mobile trong tương lai mà không cần thay đổi kiến trúc. Các dịch vụ bên ngoài (VNPay, MoMo, Gemini AI) được kết nối qua Backend, đảm bảo API key và thông tin nhạy cảm không bị lộ ra phía client.')

    # 3.2 Sequence Diagrams
    add_heading(doc, '3.2. Thiết kế Sơ đồ tuần tự (Sequence Diagram)', size=14, color=(31, 56, 100))
    add_para(doc, 'Sơ đồ tuần tự (Sequence Diagram) là loại sơ đồ UML mô tả luồng tương tác giữa các đối tượng/thành phần trong hệ thống theo trình tự thời gian (từ trên xuống dưới). Ký hiệu sử dụng trong các sơ đồ:')
    add_bullet(doc, 'Hộp chữ nhật ở đầu (Participant): đại diện cho các thành phần tham gia (User, Frontend, Backend API, Database, Admin Panel).')
    add_bullet(doc, 'Đường đứt nét dọc (Lifeline): thể hiện vòng đời hoạt động của mỗi thành phần theo thời gian.')
    add_bullet(doc, 'Hình chữ nhật mỏng trên lifeline (Activation Box): thể hiện khoảng thời gian thành phần đang xử lý.')
    add_bullet(doc, 'Mũi tên liền nét (→): thông điệp gửi đi (yêu cầu, gọi hàm).')
    add_bullet(doc, 'Mũi tên đứt nét (←): thông điệp phản hồi (kết quả trả về).')
    add_bullet(doc, 'Khung "alt" màu đỏ/xanh: biểu diễn nhánh điều kiện (alternative fragment) — xử lý trường hợp lỗi hoặc nhánh thay thế.')

    add_heading(doc, '3.2.1. Sơ đồ tuần tự – Đăng nhập hệ thống', size=13, color=(70, 100, 160))
    add_para(doc, 'Mô tả luồng xử lý: Chức năng đăng nhập là cổng vào của mọi chức năng yêu cầu xác thực. Hệ thống sử dụng Laravel Sanctum để phát hành Personal Access Token sau khi xác thực thành công. Token này được Frontend lưu vào localStorage/Cookie và gửi kèm trong header "Authorization: Bearer {token}" cho mọi request tiếp theo cần xác thực.')
    add_image(doc, 'seq_login', width_cm=12)
    add_figure_caption(doc, 'Hình 3.2 – Sơ đồ tuần tự: Đăng nhập hệ thống')
    add_steps(doc,
        [
            'Người dùng mở trang đăng nhập, nhập địa chỉ email và mật khẩu vào form.',
            'Frontend (Next.js) gửi yêu cầu HTTP POST /api/login với body {email, password} dưới dạng JSON đến Backend API.',
            'Backend API (Laravel) truy vấn cơ sở dữ liệu: SELECT * FROM users WHERE email = ? để tìm kiếm tài khoản.',
            'Database trả về bản ghi người dùng nếu tìm thấy.',
            'Backend kiểm tra mật khẩu bằng hàm bcrypt_check() — so sánh mật khẩu người dùng nhập với hash được lưu trong CSDL.',
            '[Luồng chính] Nếu đúng: Backend tạo và trả về Sanctum Personal Access Token kèm thông tin người dùng (HTTP 200 OK).',
            'Frontend nhận token, lưu vào localStorage, sau đó redirect người dùng đến trang chủ và hiển thị trạng thái đã đăng nhập.',
            '[Luồng thay thế] Nếu sai email hoặc mật khẩu: Backend trả về HTTP 401 Unauthorized không kèm token.',
            'Frontend hiển thị thông báo lỗi đỏ "Email hoặc mật khẩu không đúng" và giữ người dùng tại trang đăng nhập.',
        ],
        title='Mô tả chi tiết từng bước trong sơ đồ:'
    )

    add_heading(doc, '3.2.2. Sơ đồ tuần tự – Đặt hàng (Checkout)', size=13, color=(70, 100, 160))
    add_para(doc, 'Chức năng đặt hàng là nghiệp vụ phức tạp nhất, yêu cầu thực hiện nhiều thao tác CSDL nguyên tử trong một giao dịch duy nhất. Hệ thống hỗ trợ ba phương thức thanh toán: COD, VNPay và MoMo.')
    add_image(doc, 'seq_checkout', width_cm=12)
    add_figure_caption(doc, 'Hình 3.3 – Sơ đồ tuần tự: Đặt hàng (Checkout)')
    add_steps(doc,
        [
            'Người dùng truy cập trang Giỏ hàng, Frontend gửi GET /api/cart (kèm Bearer token) để lấy danh sách sản phẩm trong giỏ.',
            'Backend trả về danh sách cart_items (sản phẩm, số lượng, giá), Frontend hiển thị cho người dùng xem.',
            'Người dùng nhập địa chỉ giao hàng, chọn phương thức thanh toán (COD/VNPay/MoMo), nhập mã coupon (nếu có), rồi nhấn "Đặt hàng".',
            'Frontend gửi POST /api/checkout với body {address, payment_method, coupon_id} kèm Bearer token đến Backend.',
            'Backend kiểm tra tồn kho: SELECT quantity FROM products WHERE id IN (danh sách sản phẩm) để xác nhận hàng còn đủ.',
            '[Luồng chính] Backend mở giao dịch BEGIN TRANSACTION: INSERT vào bảng orders (tạo đơn hàng mới với trạng thái pending).',
            'Backend INSERT vào bảng order_items: lưu từng sản phẩm và giá tại thời điểm đặt hàng (price snapshot).',
            'Backend UPDATE bảng products: trừ số lượng tồn kho (quantity = quantity - số lượng đặt).',
            'Backend DELETE khỏi bảng cart_items: xóa giỏ hàng của người dùng sau khi đặt hàng thành công.',
            'Backend COMMIT giao dịch và trả về HTTP 201 Created kèm order_id cho Frontend.',
            'Frontend hiển thị trang xác nhận đặt hàng thành công với mã đơn hàng.',
            '[Luồng thay thế] Nếu kiểm tra tồn kho thất bại (hết hàng): Backend ROLLBACK toàn bộ giao dịch và trả về 422 Unprocessable Entity.',
            'Frontend hiển thị thông báo lỗi hết hàng, không tạo đơn hàng, tồn kho không thay đổi.',
        ],
        title='Mô tả chi tiết từng bước trong sơ đồ:'
    )

    add_heading(doc, '3.2.3. Sơ đồ tuần tự – Đánh giá sản phẩm', size=13, color=(70, 100, 160))
    add_para(doc, 'Hệ thống chỉ cho phép đánh giá khi đơn hàng có trạng thái "completed" nhằm đảm bảo tính xác thực của review và ngăn chặn đánh giá giả mạo.')
    add_image(doc, 'seq_review', width_cm=12)
    add_figure_caption(doc, 'Hình 3.4 – Sơ đồ tuần tự: Đánh giá sản phẩm')
    add_steps(doc,
        [
            'Người dùng truy cập trang Lịch sử Đơn hàng, Frontend gửi GET /api/orders kèm Bearer token.',
            'Backend truy vấn CSDL SELECT orders WHERE user_id, trả về danh sách đơn hàng.',
            'Frontend hiển thị danh sách đơn hàng kèm trạng thái. Người dùng chọn đơn hàng có trạng thái "completed" và nhấn nút "Đánh giá".',
            'Frontend hiển thị form đánh giá: chọn số sao (1–5 sao) và ô nhập nội dung nhận xét.',
            'Người dùng nhập đánh giá và nhấn Gửi. Frontend gửi POST /api/review với {product_id, rating, comment}.',
            'Backend kiểm tra điều kiện 1: SELECT orders WHERE user_id = ? AND status = completed. Xác nhận người dùng đã mua hàng.',
            'Backend kiểm tra điều kiện 2: Xác nhận product_id có tồn tại trong order_items của đơn hàng đó không.',
            '[Luồng chính] Cả hai điều kiện đều hợp lệ: Backend INSERT vào bảng reviews (user_id, product_id, rating, comment).',
            'Backend cập nhật avg_rating của sản phẩm: UPDATE products SET avg_rating = AVG(rating) WHERE id = product_id.',
            'Backend trả về HTTP 200 OK. Frontend hiển thị thông báo "Đánh giá thành công".',
            '[Luồng thay thế] Nếu đơn hàng chưa giao: Backend trả về HTTP 403 Forbidden. Frontend thông báo "Chỉ đánh giá sau khi nhận hàng".',
        ],
        title='Mô tả chi tiết từng bước trong sơ đồ:'
    )

    add_heading(doc, '3.2.4. Sơ đồ tuần tự – Admin quản lý sản phẩm', size=13, color=(70, 100, 160))
    add_para(doc, 'Admin thực hiện toàn bộ CRUD (Create – Read – Update – Delete) sản phẩm qua Admin Panel. Mọi request được bảo vệ bởi middleware auth:sanctum + admin. Hệ thống hỗ trợ upload đa ảnh cho mỗi sản phẩm.')
    add_image(doc, 'seq_admin_product', width_cm=12)
    add_figure_caption(doc, 'Hình 3.5 – Sơ đồ tuần tự: Admin quản lý sản phẩm')
    add_steps(doc,
        [
            'Admin đăng nhập Admin Panel và truy cập mô-đun Quản lý Sản phẩm.',
            'Frontend gửi GET /api/admin/products (kèm admin token), Backend trả về danh sách sản phẩm JOIN bảng categories.',
            '[THÊM SẢN PHẨM] Admin điền form: tên, mô tả, giá, danh mục, số lượng, chọn file ảnh đại diện và các ảnh bổ sung.',
            'Frontend gửi POST /api/admin/products dưới dạng multipart/form-data (để gửi cả text lẫn file ảnh).',
            'Backend INSERT vào bảng products (các thông số cơ bản + đường dẫn ảnh chính) và INSERT vào bảng product_images (các ảnh bổ sung).',
            'Backend trả về 201 Created. Admin Panel hiển thị thông báo "Thêm sản phẩm thành công" và cập nhật danh sách.',
            '[SỪ SẢN PHẨM] Admin chọn sản phẩm, chỉnh sửa thông tin và nhấn Lưu.',
            'Frontend gửi PUT /api/admin/products/{id} với dữ liệu đã cập nhật. Backend UPDATE bảng products WHERE id. Trả về 200 OK.',
            '[XÓA SẢN PHẨM] Admin xác nhận hộp thoại xóa. Frontend gửi DELETE /api/admin/products/{id}.',
            'Backend xửa bản ghi trong bảng products. Do có ON DELETE CASCADE, các bản ghi trong bảng product_images liên quan cũng bị xóa tự động.',
            'Backend trả về 200 OK. Admin Panel cập nhật danh sách, sản phẩm biến mất khỏi danh sách.',
        ],
        title='Mô tả chi tiết từng bước trong sơ đồ:'
    )

    add_heading(doc, '3.2.5. Sơ đồ tuần tự – Vòng quay may mắn (Lucky Wheel)', size=13, color=(70, 100, 160))
    add_para(doc, 'Vòng quay may mắn là tính năng gamification giới hạn 1 lần/ngày/người dùng. Phần thưởng được chọn ngẫu nhiên theo xác suất: 5% Coupon 10%, 10% Coupon 5%, 10% Coupon 50K, 75% không trúng.')
    add_image(doc, 'seq_lucky_wheel', width_cm=12)
    add_figure_caption(doc, 'Hình 3.6 – Sơ đồ tuần tự: Vòng quay may mắn (Lucky Wheel)')
    add_steps(doc,
        [
            'Người dùng truy cập trang Vòng quay may mắn và nhấn nút "QUAY".',
            'Frontend gửi POST /api/lucky-wheel/spin kèm Bearer token đến Backend.',
            'Backend kiểm tra điều kiện: SELECT * FROM lucky_wheel_spins WHERE user_id = ? AND DATE(created_at) = CURDATE().',
            '[Luồng thay thế] Nếu đã quay hôm nay: Backend trả về HTTP 429 Too Many Requests. Frontend hiển thị "Bạn đã quay hôm nay rồi, quay lại vào ngày mai!".',
            '[Luồng chính] Nếu chưa quay: Backend thực hiện random phần thưởng theo xác suất định sẵn (PHP rand / weighted random).',
            'Backend INSERT vào bảng lucky_wheel_spins: ghi lại user_id, prize_name, coupon_code, thời gian quay.',
            '[Nếu trúng coupon] Backend tự động INSERT vào bảng coupons: tạo mã ngẫu nhiên duy nhất, gán giá trị, đặt is_active = 1, valid_until = ngày hôm nay + 7 ngày.',
            'Backend trả về HTTP 200 OK với dữ liệu {prize_name, coupon_code} (coupon_code = null nếu không trúng).',
            'Frontend hiển thị animation vòng quay quay trong 3–5 giây, sau đó dừng lại ở ô phần thưởng tương ứng.',
            'Frontend hiển thị popup kết quả: nếu trúng coupon thì hiển thị mã coupon để người dùng chép lại.',
        ],
        title='Mô tả chi tiết từng bước trong sơ đồ:'
    )

    add_heading(doc, '3.2.6. Sơ đồ tuần tự – Live Chat hỗ trợ khách hàng', size=13, color=(70, 100, 160))
    add_para(doc, 'Live Chat là kênh hỗ trợ hai chiều giữa khách hàng và admin theo cơ chế short polling (3 giây/lần). Mỗi cuộc hội thoại lưu trong bảng chat_sessions, tin nhắn trong bảng chat_messages với sender_type phân biệt user/admin.')
    add_image(doc, 'seq_livechat', width_cm=14)
    add_figure_caption(doc, 'Hình 3.7 – Sơ đồ tuần tự: Live Chat hỗ trợ khách hàng')
    add_steps(doc,
        [
            'Khách hàng nhấn nút "Chat với chúng tôi" trên website. Frontend gửi POST /api/live-chat/start để tạo phiên mới.',
            'Backend INSERT vào bảng chat_sessions với guest_id ngẫu nhiên và status = active. Trả về session_id.',
            'Frontend hiển thị cửa sổ chat. Khách hàng nhập tin nhắn và gửi.',
            'Frontend gửi POST /api/live-chat/{sessionId}/messages với {message, sender_type: "user"}.',
            'Backend INSERT vào bảng chat_messages (chat_session_id, message, sender_type = user).',
            '[Phía Admin] Admin Panel định kỳ gọi GET /api/admin/live-chat/sessions (mỗi 3 giây) để phát hiện session mới và tin nhắn chưa đọc.',
            'Backend SELECT các chat_sessions có status = active kèm tin nhắn mới nhất, trả về danh sách cho Admin Panel.',
            'Admin Panel hiển thị danh sách chat đang đợi, Admin chọn phiên và gõ phản hồi.',
            'Admin Panel gửi POST /api/admin/live-chat/sessions/{id}/messages với {message, sender_type: "admin"}.',
            'Backend INSERT vào bảng chat_messages (sender_type = admin). Trả về 200 OK.',
            '[Phía User] Frontend của khách hàng định kỳ gọi GET /api/live-chat/{sessionId}/messages để lấy tin nhắn mới.',
            'Backend trả về tất cả chat_messages của session. Frontend hiển thị phản hồi của admin cho khách hàng.',
        ],
        title='Mô tả chi tiết từng bước trong sơ đồ:'
    )

    # 3.3 ERD + Data Dictionary
    add_heading(doc, '3.3. Sơ đồ Quan hệ Cơ sở dữ liệu (ERD & Data Dictionary)', size=14, color=(31, 56, 100))
    add_para(doc, 'Sơ đồ Thực thể – Quan hệ (ERD – Entity Relationship Diagram) là công cụ mô hình hóa CSDL quan hệ, thể hiện các thực thể (Entity), thuộc tính (Attribute) và quan hệ (Relationship) giữa chúng. Ký hiệu sử dụng trong ERD:')
    add_bullet(doc, 'PK (Primary Key) – màu vàng: khóa chính, định danh duy nhất mỗi bản ghi trong bảng.')
    add_bullet(doc, 'FK (Foreign Key) – màu xanh: khóa ngoại, tham chiếu đến khóa chính của bảng khác, thiết lập quan hệ giữa các thực thể.')
    add_bullet(doc, 'Ký hiệu 1:N (One-to-Many): một bản ghi ở bảng A có thể liên kết với nhiều bản ghi ở bảng B (ví dụ: 1 users có N orders).')
    add_bullet(doc, 'ON DELETE CASCADE: khi xóa bản ghi cha, tất cả bản ghi con liên quan cũng bị xóa tự động — áp dụng cho các bảng order_items, product_images, cart_items.')
    add_para(doc, 'Hệ thống sử dụng CSDL MySQL với 15 bảng, được chia thành 2 nhóm: bảng cốt lõi (8 bảng) và bảng tính năng mở rộng (7 bảng):')

    add_heading(doc, '3.3.1. ERD – Các bảng chính', size=13, color=(70, 100, 160))
    add_para(doc, 'Sơ đồ ERD phần 1 thể hiện 8 bảng cốt lõi của hệ thống thương mại điện tử, bao gồm: users (người dùng), categories (danh mục), products (sản phẩm), product_images (ảnh đa ảnh), orders (đơn hàng), order_items (chi tiết đơn hàng), cart_items (giỏ hàng) và reviews (đánh giá). Các bảng này tạo thành xương sống của toàn bộ nghiệp vụ mua bán.')
    add_image(doc, 'erd_main', width_cm=15.5)
    add_figure_caption(doc, 'Hình 3.8 – Sơ đồ ERD (Phần 1): Quan hệ các bảng chính')
    add_para(doc, 'Phân tích sơ đồ ERD phần 1: Bảng users là trung tâm, liên kết 1:N với orders, cart_items và reviews. Bảng products cũng là thực thể trung tâm quan trọng, liên kết với categories (N:1), product_images (1:N), order_items (1:N), cart_items (1:N) và reviews (1:N). Bảng order_items đóng vai trò bảng trung gian (Junction Table) giữa orders và products, giải quyết quan hệ N:N — một đơn hàng có nhiều sản phẩm, một sản phẩm có thể xuất hiện trong nhiều đơn hàng. Giá (price) được lưu lại tại thời điểm đặt hàng (snapshot), đảm bảo lịch sử đơn hàng chính xác dù giá sản phẩm sau này thay đổi.')

    add_heading(doc, '3.3.2. ERD – Các bảng tính năng nâng cao', size=13, color=(70, 100, 160))
    add_para(doc, 'Sơ đồ ERD phần 2 thể hiện 7 bảng cho các tính năng mở rộng và cấu hình hệ thống: coupons (mã giảm giá), lucky_wheel_spins (lịch sử vòng quay), chat_sessions (phiên live chat), chat_messages (tin nhắn chat), contacts (yêu cầu liên hệ), payment_settings (cấu hình thanh toán) và homepage_settings (cài đặt trang chủ).')
    add_image(doc, 'erd_extended', width_cm=15.5)
    add_figure_caption(doc, 'Hình 3.9 – Sơ đồ ERD (Phần 2): Coupons, Chat, Contacts, Settings')
    add_para(doc, 'Phân tích sơ đồ ERD phần 2: Bảng coupons liên kết với orders qua FK coupon_id (nullable) — một đơn hàng có thể áp dụng tối đa 1 coupon, nhưng một coupon có thể được nhiều đơn hàng sử dụng (theo usage_limit). Bảng lucky_wheel_spins liên kết với users (FK user_id) để theo dõi lịch sử quay theo người dùng và kiểm soát giới hạn 1 lần/ngày. Bảng chat_sessions và chat_messages có quan hệ 1:N, cho phép một phiên chat chứa nhiều tin nhắn từ cả user lẫn admin. Bảng payment_settings và homepage_settings theo mô hình Key-Value Store — linh hoạt lưu trữ các cấu hình động mà không cần thêm cột mới khi bổ sung cấu hình.')

    # 3.3.3 Data Dictionary
    add_heading(doc, '3.3.3. Từ điển dữ liệu (Data Dictionary)', size=13, color=(70, 100, 160))
    add_para(doc, 'Dưới đây là mô tả chi tiết cấu trúc từng bảng trong cơ sở dữ liệu:')

    tables_info = [
        {
            'name': 'Bảng users – Thông tin người dùng',
            'fig': 'Bảng 3.1',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính tự tăng'),
                ('name', 'VARCHAR(100)', 'NOT NULL', 'Tên đầy đủ của người dùng'),
                ('email', 'VARCHAR(100)', 'NOT NULL, UNIQUE', 'Email đăng nhập (duy nhất)'),
                ('password', 'VARCHAR(255)', 'NOT NULL', 'Mật khẩu mã hóa bcrypt'),
                ('phone', 'VARCHAR(20)', 'NULL', 'Số điện thoại liên hệ'),
                ('address', 'TEXT', 'NULL', 'Địa chỉ nhận hàng mặc định'),
                ('role', "ENUM('user','admin')", "DEFAULT 'user'", 'Vai trò trong hệ thống'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm tạo / cập nhật'),
            ]
        },
        {
            'name': 'Bảng categories – Danh mục sản phẩm',
            'fig': 'Bảng 3.2',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính tự tăng'),
                ('name', 'VARCHAR(100)', 'NOT NULL', 'Tên danh mục (Gaming, Văn phòng...)'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm tạo / cập nhật'),
            ]
        },
        {
            'name': 'Bảng products – Thông tin sản phẩm laptop',
            'fig': 'Bảng 3.3',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
                ('name', 'VARCHAR(200)', 'NOT NULL', 'Tên sản phẩm laptop'),
                ('description', 'TEXT', 'NOT NULL', 'Mô tả chi tiết thông số kỹ thuật'),
                ('price', 'DECIMAL(15,2)', 'NOT NULL', 'Giá gốc sản phẩm (VND)'),
                ('discount', 'INT', 'NULL', 'Phần trăm giảm giá (0–100)'),
                ('quantity', 'INT', 'NOT NULL', 'Số lượng tồn kho'),
                ('category_id', 'BIGINT UNSIGNED', 'FK → categories', 'Khóa ngoại danh mục'),
                ('image', 'VARCHAR(255)', 'NOT NULL', 'Đường dẫn ảnh đại diện chính'),
                ('avg_rating', 'DECIMAL(3,2)', 'NULL', 'Điểm đánh giá trung bình (1.00–5.00)'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Ngày thêm / cập nhật'),
            ]
        },
        {
            'name': 'Bảng product_images – Hình ảnh đa ảnh sản phẩm',
            'fig': 'Bảng 3.4',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
                ('product_id', 'BIGINT UNSIGNED', 'FK → products', 'Sản phẩm liên kết (ON DELETE CASCADE)'),
                ('image_url', 'VARCHAR(500)', 'NOT NULL', 'URL ảnh bổ sung của sản phẩm'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm upload'),
            ]
        },
        {
            'name': 'Bảng orders – Đơn hàng',
            'fig': 'Bảng 3.5',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3.5), Cm(4), Cm(2.5), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('user_id', 'BIGINT UNSIGNED', 'FK → users', 'Người dùng đặt hàng'),
                ('total_price', 'DECIMAL(15,2)', 'NOT NULL', 'Tổng tiền đơn hàng sau giảm giá'),
                ('payment_method', "ENUM('COD','bank_transfer','vnpay','momo')", 'NOT NULL', 'Phương thức thanh toán'),
                ('status', "ENUM 5 trạng thái", "DEFAULT 'pending'", 'pending/processing/shipping/completed/cancelled'),
                ('coupon_id', 'BIGINT UNSIGNED', 'FK → coupons, NULL', 'Coupon đã áp dụng (nếu có)'),
                ('discount_amount', 'DECIMAL(15,2)', "DEFAULT 0", 'Số tiền được giảm bởi coupon'),
                ('note', 'TEXT', 'NULL', 'Ghi chú từ người dùng'),
                ('created_at / updated_at', 'TIMESTAMP', 'NOT NULL', 'Ngày đặt / cập nhật trạng thái'),
            ]
        },
        {
            'name': 'Bảng order_items – Chi tiết đơn hàng',
            'fig': 'Bảng 3.6',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3.5), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('order_id', 'BIGINT UNSIGNED', 'FK → orders', 'Đơn hàng chứa sản phẩm'),
                ('product_id', 'BIGINT UNSIGNED', 'FK → products', 'Sản phẩm được mua'),
                ('quantity', 'INT', 'NOT NULL', 'Số lượng đặt mua'),
                ('price', 'DECIMAL(15,2)', 'NOT NULL', 'Giá tại thời điểm đặt hàng (snapshot)'),
            ]
        },
        {
            'name': 'Bảng cart_items – Giỏ hàng',
            'fig': 'Bảng 3.7',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3.5), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('user_id', 'BIGINT UNSIGNED', 'FK → users', 'Người dùng sở hữu giỏ hàng'),
                ('product_id', 'BIGINT UNSIGNED', 'FK → products', 'Sản phẩm trong giỏ'),
                ('quantity', 'INT', 'NOT NULL', 'Số lượng sản phẩm trong giỏ'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Ngày thêm / cập nhật số lượng'),
            ]
        },
        {
            'name': 'Bảng reviews – Đánh giá sản phẩm',
            'fig': 'Bảng 3.8',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('user_id', 'BIGINT UNSIGNED', 'FK → users', 'Người dùng đánh giá'),
                ('product_id', 'BIGINT UNSIGNED', 'FK → products', 'Sản phẩm được đánh giá'),
                ('rating', 'INT', 'NOT NULL, CHECK 1–5', 'Số sao đánh giá (1 đến 5)'),
                ('comment', 'TEXT', 'NOT NULL', 'Nội dung nhận xét chi tiết'),
                ('created_at', 'TIMESTAMP', 'NOT NULL', 'Ngày viết đánh giá'),
            ]
        },
        {
            'name': 'Bảng coupons – Mã giảm giá',
            'fig': 'Bảng 3.9',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3.5), Cm(3.5), Cm(3), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('code', 'VARCHAR(255)', 'NOT NULL, UNIQUE', 'Mã coupon duy nhất'),
                ('type', "ENUM('percent','fixed')", 'NOT NULL', 'Loại: phần trăm hoặc cố định'),
                ('value', 'DECIMAL(15,2)', 'NOT NULL', 'Giá trị giảm giá'),
                ('min_order_value', 'DECIMAL(15,2)', "DEFAULT 0", 'Giá trị đơn hàng tối thiểu'),
                ('max_discount', 'DECIMAL(15,2)', 'NULL', 'Mức giảm tối đa (cho percent)'),
                ('usage_limit / used_count', 'INT', 'NULL / DEFAULT 0', 'Giới hạn / số lần đã dùng'),
                ('valid_from / valid_until', 'TIMESTAMP', 'NULL', 'Thời gian hiệu lực của coupon'),
                ('is_active', 'TINYINT(1)', "DEFAULT 1", 'Trạng thái kích hoạt coupon'),
            ]
        },
        {
            'name': 'Bảng lucky_wheel_spins – Lịch sử quay may mắn',
            'fig': 'Bảng 3.10',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3.5), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('user_id', 'BIGINT UNSIGNED', 'FK → users', 'Người dùng quay'),
                ('prize_name', 'VARCHAR(255)', 'NULL', 'Tên phần thưởng nhận được'),
                ('coupon_code', 'VARCHAR(255)', 'NULL', 'Mã coupon trúng thưởng (nếu có)'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm quay'),
            ]
        },
        {
            'name': 'Bảng chat_sessions – Phiên hội thoại live chat',
            'fig': 'Bảng 3.11',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3.5), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('guest_id', 'VARCHAR(255)', 'NOT NULL', 'Session ID ẩn danh của khách'),
                ('status', 'VARCHAR(255)', "DEFAULT 'active'", 'Trạng thái: active / closed'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm mở / cập nhật session'),
            ]
        },
        {
            'name': 'Bảng chat_messages – Tin nhắn live chat',
            'fig': 'Bảng 3.12',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3.5), Cm(3.5), Cm(3), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('chat_session_id', 'BIGINT UNSIGNED', 'FK → chat_sessions', 'Phiên hội thoại chứa tin nhắn'),
                ('sender_type', 'VARCHAR(255)', "NOT NULL", "Người gửi: 'user' hoặc 'admin'"),
                ('message', 'TEXT', 'NOT NULL', 'Nội dung tin nhắn'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm gửi tin nhắn'),
            ]
        },
        {
            'name': 'Bảng contacts – Yêu cầu liên hệ',
            'fig': 'Bảng 3.13',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3.5), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('name', 'VARCHAR(255)', 'NOT NULL', 'Tên người liên hệ'),
                ('email / phone', 'VARCHAR(255)', 'NULL', 'Email / SĐT liên hệ'),
                ('message', 'TEXT', 'NOT NULL', 'Nội dung yêu cầu hỏi đáp'),
                ('status', "ENUM('new','processing','resolved')", "DEFAULT 'new'", 'Trạng thái xử lý yêu cầu'),
                ('admin_note', 'TEXT', 'NULL', 'Ghi chú của admin khi xử lý'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Ngày gửi / cập nhật'),
            ]
        },
        {
            'name': 'Bảng payment_settings – Cấu hình thanh toán',
            'fig': 'Bảng 3.14',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3.5), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('setting_key', 'VARCHAR(120)', 'NOT NULL, UNIQUE', 'Tên key (vnpay_url, momo_key...)'),
                ('setting_value', 'TEXT', 'NULL', 'Giá trị cấu hình tương ứng'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Ngày tạo / cập nhật cấu hình'),
            ]
        },
        {
            'name': 'Bảng homepage_settings – Cài đặt giao diện trang chủ',
            'fig': 'Bảng 3.15',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3.5), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK', 'Khóa chính'),
                ('setting_key', 'VARCHAR(120)', 'NOT NULL, UNIQUE', 'Tên key (banner_url, slider_data...)'),
                ('setting_value', 'LONGTEXT', 'NULL', 'Giá trị (có thể là JSON dài)'),
                ('created_at / updated_at', 'TIMESTAMP', 'NULL', 'Ngày tạo / cập nhật cài đặt'),
            ]
        },
    ]

    for tbl_info in tables_info:
        add_heading(doc, tbl_info['name'], size=12, color=(31, 56, 100), space_before=10, space_after=4)
        ncols = len(tbl_info['cols'])
        table = doc.add_table(rows=len(tbl_info['rows'])+1, cols=ncols)
        table.style = 'Table Grid'
        set_table_borders(table)
        add_table_header_row(table, tbl_info['cols'])
        for i, row_data in enumerate(tbl_info['rows']):
            row = table.rows[i+1]
            for j, cell_val in enumerate(row_data):
                row.cells[j].text = ''
                p = row.cells[j].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(cell_val)
                set_font(r, bold=(j == 0), size=10)
                if j == 2 and 'PK' in cell_val:
                    r.font.color.rgb = RGBColor(180, 20, 20)
                elif j == 2 and 'FK' in cell_val:
                    r.font.color.rgb = RGBColor(0, 80, 160)
            if i % 2 == 1:
                for j in range(ncols):
                    tcPr = table.rows[i+1].cells[j]._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'F5F8FF')
                    tcPr.append(shd)
        for j, w in enumerate(tbl_info['widths']):
            for row in table.rows:
                row.cells[j].width = w
        add_figure_caption(doc, f"{tbl_info['fig']} – Cấu trúc {tbl_info['name'].split('–')[0].strip()}")
        doc.add_paragraph()

    # 3.4 Công nghệ
    add_heading(doc, '3.4. Cơ sở lý thuyết về Công nghệ cốt lõi', size=14, color=(31, 56, 100))
    tech_data = [
        ('Laravel 10 (PHP 8.x)', 'Framework PHP theo mô hình MVC. Cung cấp Eloquent ORM, Sanctum authentication, middleware, RESTful routing. Xây dựng toàn bộ Backend API.'),
        ('Next.js 13+ (React)', 'Framework React hỗ trợ SSR/SSG. Dùng cho Frontend người dùng và Admin Panel. Routing động, API routes.'),
        ('MySQL 8.x', 'RDBMS lưu trữ toàn bộ dữ liệu với ràng buộc FK, UNIQUE, NOT NULL đảm bảo toàn vẹn dữ liệu.'),
        ('Laravel Sanctum', 'Package xác thực API token-based. Phát hành token sau đăng nhập, middleware auth:sanctum bảo vệ routes.'),
        ('React Context API', 'Quản lý trạng thái toàn cục phía frontend (user session, giỏ hàng, thông báo).'),
        ('VNPay / MoMo API', 'Cổng thanh toán điện tử phổ biến tại Việt Nam. Redirect sang cổng thanh toán và xử lý callback/IPN.'),
        ('Google Gemini API', 'AI Chatbot tư vấn sản phẩm tự động. Backend gọi Gemini với context sản phẩm thực tế của cửa hàng.'),
        ('RESTful API', 'Chuẩn thiết kế API với GET/POST/PUT/DELETE, trả về JSON. Dễ dàng tích hợp mobile app và đối tác.'),
    ]
    table = doc.add_table(rows=len(tech_data)+1, cols=2)
    table.style = 'Table Grid'
    set_table_borders(table)
    add_table_header_row(table, ['Công nghệ', 'Mô tả & Ứng dụng trong dự án'])
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(11.5)
    for i, (tech, desc) in enumerate(tech_data):
        row = table.rows[i+1]
        row.cells[0].text = ''
        row.cells[1].text = ''
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(tech)
        set_font(r0, bold=True, size=10)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(desc)
        set_font(r1, size=10)
        if i % 2 == 1:
            for j in range(2):
                tcPr = row.cells[j]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F5F8FF')
                tcPr.append(shd)
    add_figure_caption(doc, 'Bảng 3.16 – Công nghệ cốt lõi và ứng dụng trong hệ thống')

    # 3.5 DFD
    add_heading(doc, '3.5. Sơ đồ luồng dữ liệu (Data Flow Diagram – DFD)', size=14, color=(31, 56, 100))
    add_para(doc, 'Sơ đồ luồng dữ liệu (DFD – Data Flow Diagram) mô tả cách dữ liệu di chuyển qua hệ thống — từ nguồn đầu vào, qua các quy trình xử lý, đến nơi lưu trữ và đầu ra. DFD sử dụng 4 ký hiệu chuẩn theo ký pháp Gane-Sarson:')
    add_bullet(doc, 'Hình chữ nhật (External Entity): thực thể bên ngoài cung cấp hoặc nhận dữ liệu từ hệ thống — trong đề tài gồm: Người Dùng, Quản Trị Viên, VNPay/MoMo, Google Gemini AI.')
    add_bullet(doc, 'Hình tròn/oval (Process): quy trình xử lý dữ liệu bên trong hệ thống — 6 quy trình chính: Xác thực & Phân quyền, Quản lý Sản phẩm, Quản lý Giỏ hàng & Đặt hàng, Thanh toán, Đánh giá & Khuyến mãi, Hỗ trợ Khách hàng.')
    add_bullet(doc, 'Hình chữ nhật mở (Data Store): kho lưu trữ dữ liệu — tương ứng với các bảng CSDL: D1:users, D2:products, D3:orders, D4:coupons, D5:reviews, D6:chat_messages.')
    add_bullet(doc, 'Mũi tên có nhãn (Data Flow): luồng dữ liệu di chuyển giữa các thành phần, nhãn mô tả loại dữ liệu được truyền.')
    add_image(doc, 'dfd', width_cm=15.5)
    add_figure_caption(doc, 'Hình 3.10 – Sơ đồ luồng dữ liệu (DFD) mức 0 và mức 1')
    add_para(doc, 'Phân tích sơ đồ DFD mức 0 (Context Diagram): Hệ thống nhận dữ liệu đầu vào từ 4 nguồn bên ngoài. Người Dùng cung cấp thông tin đăng ký, đơn hàng, đánh giá; nhận lại sản phẩm, xác nhận đơn và phản hồi. Quản Trị Viên cung cấp dữ liệu quản lý; nhận lại thống kê và báo cáo. VNPay/MoMo là cổng thanh toán hai chiều — hệ thống gửi yêu cầu và nhận kết quả. Google Gemini AI nhận câu hỏi và trả về câu trả lời tư vấn sản phẩm.')
    add_para(doc, 'Phân tích sơ đồ DFD mức 1: 6 quy trình xử lý chính được kết nối với các kho dữ liệu tương ứng. Quy trình "Xác thực & Phân quyền" đọc/ghi vào D1:users. Quy trình "Quản lý Sản phẩm" đọc/ghi vào D2:products. Quy trình "Quản lý Giỏ hàng & Đặt hàng" tương tác với D3:orders và D2:products. Quy trình "Thanh toán" đọc D4:coupons và cập nhật D3:orders. Quy trình "Đánh giá & Khuyến mãi" đọc/ghi D5:reviews và D4:coupons. Quy trình "Hỗ trợ Khách hàng" đọc/ghi D6:chat_messages.')

    # 3.6 Module khó
    add_heading(doc, '3.6. Giải pháp triển khai một số Module khó', size=14, color=(31, 56, 100))

    add_heading(doc, '3.6.1. Module Thanh toán VNPay và MoMo', size=13, color=(70, 100, 160))
    add_para(doc, 'Hệ thống tích hợp hai cổng thanh toán phổ biến. Luồng xử lý: Backend tạo URL thanh toán có ký chữ ký (HMAC-SHA512 cho VNPay), Frontend redirect sang cổng thanh toán, người dùng hoàn tất thanh toán, cổng gọi callback/IPN URL, Backend xác thực chữ ký và cập nhật trạng thái đơn hàng.')
    add_para(doc, 'Điểm khó: Cần xử lý race condition khi IPN và return URL đến gần như đồng thời. Giải pháp: Sử dụng database transaction kết hợp với kiểm tra trạng thái đơn hàng trước khi cập nhật.')

    add_heading(doc, '3.6.2. Module AI Chatbot (Google Gemini API)', size=13, color=(70, 100, 160))
    add_para(doc, 'Hệ thống tích hợp Google Gemini API để tự động tư vấn sản phẩm. Khi người dùng đặt câu hỏi, Backend: (1) lấy danh sách sản phẩm thực tế từ DB, (2) xây dựng prompt bao gồm context sản phẩm + câu hỏi người dùng, (3) gọi Gemini API, (4) trả kết quả về Frontend. Giải pháp này đảm bảo AI luôn tư vấn dựa trên sản phẩm thực tế trong kho.')

    add_heading(doc, '3.6.3. Module Vòng quay may mắn (Lucky Wheel)', size=13, color=(70, 100, 160))
    add_para(doc, 'Module gamification giới hạn mỗi user chỉ được quay 1 lần/ngày. Cơ chế phân bổ phần thưởng theo xác suất: 5% Coupon 10%, 10% Coupon 5%, 10% Coupon 50K cố định, 75% Không trúng. Khi trúng coupon, hệ thống tự động tạo coupon mới với code ngẫu nhiên và thời hạn 7 ngày, lưu vào bảng coupons và lucky_wheel_spins.')

    add_heading(doc, '3.6.4. Module Live Chat theo thời gian thực', size=13, color=(70, 100, 160))
    add_para(doc, 'Live chat hoạt động theo cơ chế polling (không dùng WebSocket để đơn giản hóa). Client polling mỗi 3 giây để lấy tin nhắn mới. Dữ liệu lưu trong bảng chat_sessions (quản lý phiên) và chat_messages (tin nhắn). Admin panel có thể xem tất cả session đang active, chọn session để trả lời và đóng session khi hoàn tất hỗ trợ.')


# ═══════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════

def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    build_chapter2(doc)
    build_chapter3(doc)

    out_path = '/Users/nguyennghia/Desktop/workspace/ManhLaptop/Chuong2_Chuong3_CoHinh.docx'
    doc.save(out_path)
    print(f'\n✅ File Word với hình ảnh đã được tạo: {out_path}')
    return out_path


if __name__ == '__main__':
    main()
