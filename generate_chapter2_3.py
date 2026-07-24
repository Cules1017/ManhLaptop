#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo file Word Chương 2 (Đặc tả Yêu cầu) và Chương 3 (Thiết kế & Cài đặt giải pháp)
cho Báo cáo Khóa luận: Xây dựng Hệ thống Bán Laptop Trực Tuyến với Laravel và Next.js
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────── Helpers ───────────────────────────

def set_font(run, bold=False, size=12, color=None, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force Unicode font for Vietnamese
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)


def add_heading(doc, text, level=1, size=14, bold=True, color=None, space_before=12, space_after=6):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    set_font(run, bold=bold, size=size, color=color)
    return para


def add_para(doc, text, indent=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, italic=False, space_before=2, space_after=4):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if indent:
        para.paragraph_format.left_indent = Pt(indent)
        para.paragraph_format.first_line_indent = Pt(0)
    run = para.add_run(text)
    set_font(run, bold=bold, size=size, italic=italic)
    return para


def add_bullet(doc, text, level=0, size=12):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Pt(18 + level * 18)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    set_font(run, size=size)
    return para


def set_table_borders(table):
    """Add borders to table cells."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '6')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_table_header_row(table, headers, bg_color='1F3864'):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        set_font(run, bold=True, size=11, color=(255, 255, 255))
        # Set background color
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color)
        tcPr.append(shd)


def add_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11, center=False):
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else align
    run = para.add_run(text)
    set_font(run, bold=bold, size=size)


def add_figure_caption(doc, text, size=11):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    set_font(run, bold=True, size=size, italic=True, color=(80, 80, 80))
    return para


def add_page_break(doc):
    doc.add_page_break()


# ─────────────────────────── ASCII / Text Diagrams ───────────────────────────

def add_ascii_box(doc, lines, title=None, font_name="Courier New", size=9):
    """Draw a box using table with monospace font for ASCII art diagrams."""
    if title:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        set_font(run, bold=True, size=10)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    # Cell padding
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'left', 'bottom', 'right']:
        mar = OxmlElement(f'w:{side}')
        mar.set(qn('w:w'), '80')
        mar.set(qn('w:type'), 'dxa')
        tcMar.append(mar)
    tcPr.append(tcMar)
    # Border
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), '1F3864')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    # Shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F4FA')
    tcPr.append(shd)

    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(size)

    return table


# ═══════════════════════════════════════════════════════════════
#  CHAPTER 2 – ĐẶC TẢ YÊU CẦU
# ═══════════════════════════════════════════════════════════════

def build_chapter2(doc):
    # ── Tiêu đề chương ──
    add_heading(doc, 'CHƯƠNG 2: ĐẶC TẢ YÊU CẦU', level=1, size=16, color=(31, 56, 100), space_before=18, space_after=12)

    # ─── 2.1 Khảo sát ───
    add_heading(doc, '2.1. Khảo sát hiện trạng và bài toán đặt ra', level=2, size=14, color=(31, 56, 100))
    add_para(doc, 'Trong bối cảnh thương mại điện tử phát triển mạnh mẽ, nhu cầu mua sắm laptop trực tuyến ngày càng tăng cao. Nhiều website bán laptop hiện nay vẫn còn hạn chế về tính năng quản lý, trải nghiệm người dùng và bảo mật. Đề tài hướng đến xây dựng một hệ thống bán laptop trực tuyến hoàn chỉnh với đầy đủ chức năng hiện đại, đáp ứng nhu cầu thực tế của người mua lẫn quản trị viên.')
    add_para(doc, 'Hệ thống được xây dựng gồm hai phần chính: giao diện người dùng (Frontend sử dụng Next.js/React) và hệ thống quản trị (Admin Panel), kết hợp với backend Laravel xử lý toàn bộ nghiệp vụ và API RESTful. Cơ sở dữ liệu MySQL lưu trữ toàn bộ thông tin sản phẩm, đơn hàng, người dùng, đánh giá và các tính năng mở rộng.')

    # ─── 2.2 Phân tích yêu cầu chức năng ───
    add_heading(doc, '2.2. Phân tích yêu cầu chức năng', level=2, size=14, color=(31, 56, 100))
    add_para(doc, 'Hệ thống được phân tích thành các nhóm chức năng chính theo từng đối tượng sử dụng:')

    add_heading(doc, '2.2.1. Người dùng chưa đăng nhập (Guest)', level=3, size=13, color=(70, 100, 160))
    items = [
        'Xem danh sách sản phẩm: duyệt tất cả sản phẩm, lọc theo danh mục, sắp xếp theo giá/đánh giá.',
        'Tìm kiếm sản phẩm: tìm theo từ khóa tên, mô tả sản phẩm.',
        'Xem chi tiết sản phẩm: hình ảnh lớn, mô tả đầy đủ, tồn kho, đánh giá từ người mua.',
        'Thêm vào danh sách yêu thích: lưu sản phẩm vào localStorage mà không cần đăng nhập.',
        'Đăng ký tài khoản: nhập thông tin cá nhân để tạo tài khoản mới.',
        'Đăng nhập: xác thực bằng email/mật khẩu và nhận token Sanctum.',
        'Gửi liên hệ/hỏi đáp: qua form contact hoặc live chat với admin.',
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, '2.2.2. Người dùng đã đăng nhập (User)', level=3, size=13, color=(70, 100, 160))
    items2 = [
        'Tất cả chức năng của Guest.',
        'Quản lý giỏ hàng: thêm, cập nhật số lượng, xóa sản phẩm khỏi giỏ hàng.',
        'Đặt hàng: chọn địa chỉ nhận hàng, phương thức thanh toán (COD, VNPay, MoMo), nhập mã giảm giá.',
        'Vòng quay may mắn: quay để nhận coupon giảm giá.',
        'Xem lịch sử đơn hàng: danh sách đơn, trạng thái từng đơn (pending, processing, shipping, completed, cancelled).',
        'Đánh giá sản phẩm: chỉ cho phép đánh giá sau khi đơn hàng đã giao thành công (status = completed).',
        'Quản lý tài khoản cá nhân: cập nhật tên, số điện thoại, địa chỉ; đổi mật khẩu.',
        'Live chat hỗ trợ: gửi tin nhắn trực tiếp với admin/tư vấn viên.',
    ]
    for item in items2:
        add_bullet(doc, item)

    add_heading(doc, '2.2.3. Quản trị viên (Admin)', level=3, size=13, color=(70, 100, 160))
    items3 = [
        'Đăng nhập Admin: truy cập trang quản trị riêng biệt với phân quyền admin.',
        'Quản lý sản phẩm: thêm, sửa, xóa sản phẩm; upload nhiều ảnh; quản lý tồn kho, danh mục.',
        'Quản lý danh mục: thêm, sửa, xóa danh mục sản phẩm.',
        'Quản lý đơn hàng: xem danh sách, chi tiết từng đơn, cập nhật trạng thái, hủy đơn hàng.',
        'Quản lý người dùng: xem danh sách, phân quyền user/admin, xóa tài khoản.',
        'Thống kê & báo cáo: doanh thu theo ngày/tháng, top sản phẩm bán chạy, số đơn hàng theo trạng thái, xuất PDF.',
        'Quản lý coupon: tạo, sửa, xóa mã giảm giá (loại phần trăm hoặc cố định); giới hạn lượt dùng, thời hạn sử dụng.',
        'Quản lý live chat: xem các session chat, trả lời khách hàng, đóng session.',
        'Quản lý liên hệ/hỗ trợ: xem danh sách yêu cầu liên hệ, xử lý, ghi chú trạng thái.',
        'Cài đặt trang chủ: cập nhật banner, slider, nội dung trang chủ.',
        'Cài đặt thanh toán: cấu hình VNPay, MoMo.',
    ]
    for item in items3:
        add_bullet(doc, item)

    # ─── 2.3 Yêu cầu phi chức năng ───
    add_heading(doc, '2.3. Yêu cầu phi chức năng', level=2, size=14, color=(31, 56, 100))
    nfr = [
        ('Hiệu năng', 'API phản hồi dưới 500ms trong điều kiện bình thường; frontend tải trang dưới 3 giây.'),
        ('Bảo mật', 'Xác thực qua Laravel Sanctum (token-based); phân quyền admin/user rõ ràng; mật khẩu mã hóa bcrypt; bảo vệ CSRF.'),
        ('Khả năng mở rộng', 'Kiến trúc RESTful API cho phép tích hợp thêm ứng dụng mobile, đối tác thanh toán, vận chuyển.'),
        ('Giao diện & UX', 'Giao diện thân thiện, responsive trên mọi thiết bị; thông báo toast/popup; phân trang, lọc, sắp xếp sản phẩm.'),
        ('Tính toàn vẹn dữ liệu', 'Ràng buộc khóa ngoại trong CSDL; validate đầu vào phía server; rollback khi giao dịch thất bại.'),
        ('Khả năng bảo trì', 'Codebase cấu trúc MVC rõ ràng; comment đầy đủ; dễ nâng cấp tính năng.'),
    ]
    table = doc.add_table(rows=len(nfr)+1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    add_table_header_row(table, ['Yêu cầu phi chức năng', 'Mô tả chi tiết'])
    for i, (req, desc) in enumerate(nfr):
        row = table.rows[i+1]
        add_cell_text(row.cells[0], req, bold=True, size=11)
        add_cell_text(row.cells[1], desc, size=11)
    set_table_borders(table)
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(11.5)
    add_figure_caption(doc, 'Bảng 2.1 – Yêu cầu phi chức năng của hệ thống')

    # ─── 2.4 Mô hình Use Case (Tổng quát) ───
    add_heading(doc, '2.4. Mô hình Use Case (Tổng quát)', level=2, size=14, color=(31, 56, 100))
    add_para(doc, 'Sơ đồ Use Case tổng quát mô tả các tác nhân và các chức năng chính của hệ thống:')

    # Use Case diagram dạng text
    uc_lines = [
        '  ┌─────────────────────────────────────────────────────────────────────────────────┐',
        '  │                    HỆ THỐNG BÁN LAPTOP TRỰC TUYẾN                              │',
        '  │                                                                                 │',
        '  │   ┌────────────────────────────────┐  ┌──────────────────────────────────────┐ │',
        '  │   │     NGƯỜI DÙNG (USER)           │  │       QUẢN TRỊ VIÊN (ADMIN)          │ │',
        '  │   │                                 │  │                                      │ │',
        '  │   │  ○ Xem danh sách sản phẩm       │  │  ○ Quản lý sản phẩm & danh mục      │ │',
        '  │   │  ○ Tìm kiếm & lọc sản phẩm     │  │  ○ Quản lý đơn hàng                 │ │',
        '  │   │  ○ Xem chi tiết sản phẩm        │  │  ○ Quản lý người dùng               │ │',
        '  │   │  ○ Thêm vào danh sách yêu thích │  │  ○ Thống kê & báo cáo doanh thu     │ │',
        '  │   │  ○ Đăng ký / Đăng nhập          │  │  ○ Quản lý coupon giảm giá          │ │',
        '  │   │  ○ Quản lý giỏ hàng             │  │  ○ Quản lý live chat                │ │',
        '  │   │  ○ Đặt hàng (COD/VNPay/MoMo)   │  │  ○ Quản lý liên hệ khách hàng       │ │',
        '  │   │  ○ Áp dụng mã giảm giá coupon   │  │  ○ Cài đặt trang chủ & thanh toán  │ │',
        '  │   │  ○ Vòng quay may mắn            │  │  ○ Xuất báo cáo PDF                 │ │',
        '  │   │  ○ Xem lịch sử đơn hàng         │  │                                      │ │',
        '  │   │  ○ Đánh giá sản phẩm đã mua     │  │                                      │ │',
        '  │   │  ○ Cập nhật thông tin cá nhân   │  │                                      │ │',
        '  │   │  ○ Live chat hỗ trợ             │  │                                      │ │',
        '  │   └────────────────────────────────┘  └──────────────────────────────────────┘ │',
        '  │                                                                                 │',
        '  │   Actor: [User]────────────────────────────────────────────[Admin] :Actor       │',
        '  └─────────────────────────────────────────────────────────────────────────────────┘',
    ]
    add_ascii_box(doc, uc_lines, size=8)
    add_figure_caption(doc, 'Hình 2.1 – Sơ đồ Use Case tổng quát của hệ thống')

    # ─── 2.5 Sơ đồ phân rã chức năng (FDD) ───
    add_heading(doc, '2.5. Sơ đồ phân rã chức năng (Functional Decomposition Diagram)', level=2, size=14, color=(31, 56, 100))
    add_para(doc, 'Sơ đồ phân rã chức năng mô tả cấu trúc phân cấp các chức năng của hệ thống từ mức tổng quát đến chi tiết:')

    fdd_lines = [
        '  HỆ THỐNG BÁN LAPTOP TRỰC TUYẾN',
        '  │',
        '  ├─── 1. QUẢN LÝ NGƯỜI DÙNG',
        '  │         ├── 1.1 Đăng ký tài khoản',
        '  │         ├── 1.2 Đăng nhập / Đăng xuất',
        '  │         ├── 1.3 Cập nhật thông tin cá nhân',
        '  │         └── 1.4 Phân quyền (User / Admin)',
        '  │',
        '  ├─── 2. QUẢN LÝ SẢN PHẨM',
        '  │         ├── 2.1 Xem danh sách / tìm kiếm / lọc',
        '  │         ├── 2.2 Xem chi tiết sản phẩm',
        '  │         ├── 2.3 Thêm / Sửa / Xóa sản phẩm (Admin)',
        '  │         ├── 2.4 Upload nhiều ảnh sản phẩm',
        '  │         └── 2.5 Quản lý danh mục sản phẩm',
        '  │',
        '  ├─── 3. QUẢN LÝ GIỎ HÀNG & ĐẶT HÀNG',
        '  │         ├── 3.1 Thêm / Cập nhật / Xóa giỏ hàng',
        '  │         ├── 3.2 Áp dụng mã giảm giá (Coupon)',
        '  │         ├── 3.3 Đặt hàng (COD / VNPay / MoMo)',
        '  │         └── 3.4 Xem lịch sử đơn hàng',
        '  │',
        '  ├─── 4. QUẢN LÝ ĐƠN HÀNG (ADMIN)',
        '  │         ├── 4.1 Xem danh sách & chi tiết đơn hàng',
        '  │         ├── 4.2 Cập nhật trạng thái đơn',
        '  │         └── 4.3 Hủy đơn hàng',
        '  │',
        '  ├─── 5. ĐÁNH GIÁ SẢN PHẨM',
        '  │         ├── 5.1 Đánh giá sau khi nhận hàng',
        '  │         └── 5.2 Xem đánh giá của người dùng khác',
        '  │',
        '  ├─── 6. KHUYẾN MÃI & GAMIFICATION',
        '  │         ├── 6.1 Tạo / Quản lý Coupon (Admin)',
        '  │         ├── 6.2 Áp dụng Coupon khi đặt hàng',
        '  │         └── 6.3 Vòng quay may mắn nhận Coupon',
        '  │',
        '  ├─── 7. HỖ TRỢ KHÁCH HÀNG',
        '  │         ├── 7.1 Gửi form liên hệ',
        '  │         ├── 7.2 Live Chat với Admin',
        '  │         └── 7.3 AI Chatbot tư vấn sản phẩm',
        '  │',
        '  └─── 8. THỐNG KÊ & BÁO CÁO (ADMIN)',
        '            ├── 8.1 Thống kê doanh thu theo ngày/tháng',
        '            ├── 8.2 Top sản phẩm bán chạy',
        '            ├── 8.3 Số đơn hàng theo trạng thái',
        '            └── 8.4 Xuất báo cáo PDF',
    ]
    add_ascii_box(doc, fdd_lines, size=8.5)
    add_figure_caption(doc, 'Hình 2.2 – Sơ đồ phân rã chức năng (FDD) của hệ thống')

    # ─── Bảng đặc tả Use Case chi tiết ───
    add_heading(doc, '2.6. Đặc tả Use Case chi tiết', level=2, size=14, color=(31, 56, 100))
    add_para(doc, 'Dưới đây là bảng đặc tả chi tiết các Use Case quan trọng của hệ thống:')

    # UC1: Đặt hàng
    uc_specs = [
        {
            'title': 'UC-01: Đặt hàng',
            'data': [
                ('Tên Use Case', 'Đặt hàng'),
                ('Mã Use Case', 'UC-01'),
                ('Tác nhân', 'Người dùng đã đăng nhập (User)'),
                ('Mô tả', 'Người dùng đặt hàng các sản phẩm trong giỏ hàng'),
                ('Tiền điều kiện', 'User đã đăng nhập; giỏ hàng có ít nhất 1 sản phẩm'),
                ('Hậu điều kiện', 'Đơn hàng được tạo; tồn kho được cập nhật; giỏ hàng được xóa'),
                ('Dòng sự kiện chính',
                 '1. User vào trang giỏ hàng\n'
                 '2. User nhập/chọn địa chỉ nhận hàng\n'
                 '3. User chọn phương thức thanh toán (COD/VNPay/MoMo)\n'
                 '4. User áp dụng mã coupon (nếu có)\n'
                 '5. User nhấn "Đặt hàng" → Hệ thống kiểm tra tồn kho\n'
                 '6. Hệ thống tạo đơn hàng, trừ tồn kho, xóa giỏ hàng\n'
                 '7. Hiển thị xác nhận đơn hàng thành công'),
                ('Dòng sự kiện thay thế', 'Nếu sản phẩm hết hàng → báo lỗi, không tạo đơn\n'
                 'Nếu thanh toán VNPay/MoMo thất bại → hủy đơn, hoàn kho'),
            ]
        },
        {
            'title': 'UC-02: Đánh giá sản phẩm',
            'data': [
                ('Tên Use Case', 'Đánh giá sản phẩm'),
                ('Mã Use Case', 'UC-02'),
                ('Tác nhân', 'Người dùng đã đăng nhập (User)'),
                ('Mô tả', 'Người dùng đánh giá sản phẩm đã mua sau khi nhận hàng thành công'),
                ('Tiền điều kiện', 'User đã đăng nhập; có đơn hàng trạng thái "completed"'),
                ('Hậu điều kiện', 'Đánh giá được lưu vào CSDL; rating trung bình sản phẩm được cập nhật'),
                ('Dòng sự kiện chính',
                 '1. User vào trang lịch sử đơn hàng\n'
                 '2. User chọn đơn hàng đã giao → nhấn "Đánh giá"\n'
                 '3. Hệ thống hiển thị form đánh giá (số sao 1–5, nhận xét)\n'
                 '4. User nhập đánh giá và gửi\n'
                 '5. Hệ thống kiểm tra tính hợp lệ → lưu đánh giá\n'
                 '6. Hiển thị thông báo đánh giá thành công'),
                ('Dòng sự kiện thay thế', 'Nếu đơn chưa giao → từ chối, hiển thị lỗi\n'
                 'Nếu đã đánh giá rồi → hiển thị đánh giá cũ, cho phép chỉnh sửa'),
            ]
        },
        {
            'title': 'UC-03: Quản lý đơn hàng (Admin)',
            'data': [
                ('Tên Use Case', 'Quản lý đơn hàng'),
                ('Mã Use Case', 'UC-03'),
                ('Tác nhân', 'Quản trị viên (Admin)'),
                ('Mô tả', 'Admin xem, cập nhật trạng thái và hủy đơn hàng'),
                ('Tiền điều kiện', 'Admin đã đăng nhập'),
                ('Hậu điều kiện', 'Trạng thái đơn hàng được cập nhật trong CSDL'),
                ('Dòng sự kiện chính',
                 '1. Admin truy cập trang Quản lý Đơn hàng\n'
                 '2. Hệ thống hiển thị danh sách đơn hàng\n'
                 '3. Admin lọc theo trạng thái / tìm kiếm đơn\n'
                 '4. Admin chọn đơn → cập nhật trạng thái mới\n'
                 '5. Hệ thống lưu thay đổi → hiển thị cập nhật'),
                ('Dòng sự kiện thay thế', 'Nếu đơn đã "completed" hoặc "cancelled" → không cho sửa\n'
                 'Admin có thể hủy đơn ở trạng thái pending/processing'),
            ]
        },
    ]

    for uc in uc_specs:
        add_heading(doc, uc['title'], level=3, size=12, color=(70, 100, 160), space_before=10, space_after=4)
        table = doc.add_table(rows=len(uc['data']), cols=2)
        table.style = 'Table Grid'
        set_table_borders(table)
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(12)
        for i, (field, value) in enumerate(uc['data']):
            row = table.rows[i]
            row.cells[0].text = ''
            row.cells[1].text = ''
            # Field name
            tcPr = row.cells[0]._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'DCE6F1')
            tcPr.append(shd)
            p0 = row.cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r0 = p0.add_run(field)
            set_font(r0, bold=True, size=10)
            # Value
            p1 = row.cells[1].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r1 = p1.add_run(value)
            set_font(r1, size=10)
        doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
#  CHAPTER 3 – THIẾT KẾ & CÀI ĐẶT GIẢI PHÁP
# ═══════════════════════════════════════════════════════════════

def build_chapter3(doc):
    add_page_break(doc)

    # ── Tiêu đề chương ──
    add_heading(doc, 'CHƯƠNG 3: THIẾT KẾ VÀ CÀI ĐẶT GIẢI PHÁP', level=1, size=16, color=(31, 56, 100), space_before=18, space_after=12)

    # ─── 3.1 Kiến trúc tổng thể ───
    add_heading(doc, '3.1. Thiết kế kiến trúc tổng thể (System Architecture)', level=2, size=14, color=(31, 56, 100))
    add_para(doc, 'Hệ thống được thiết kế theo mô hình kiến trúc 3 tầng kết hợp với kiến trúc microservices nhẹ, bao gồm:')
    add_bullet(doc, 'Presentation Layer: Giao diện người dùng (Next.js/React) và giao diện quản trị (Admin Next.js). Hai ứng dụng này giao tiếp với backend qua RESTful API.')
    add_bullet(doc, 'Business Logic Layer: Backend Laravel (PHP) xử lý toàn bộ nghiệp vụ, xác thực, phân quyền, tích hợp cổng thanh toán, AI chatbot (Gemini API).')
    add_bullet(doc, 'Data Layer: Cơ sở dữ liệu MySQL lưu trữ toàn bộ dữ liệu. Laravel Eloquent ORM quản lý truy vấn và quan hệ dữ liệu.')

    arch_lines = [
        '  ┌─────────────────────────────────────────────────────────────────────────────────┐',
        '  │                        TẦNG TRÌNH BÀY (Presentation Layer)                     │',
        '  │  ┌──────────────────────────────┐   ┌──────────────────────────────────────┐   │',
        '  │  │   Frontend (Next.js/React)    │   │    Admin Panel (Next.js/React)        │   │',
        '  │  │   - Trang sản phẩm            │   │    - Dashboard thống kê               │   │',
        '  │  │   - Giỏ hàng & Đặt hàng       │   │    - Quản lý sản phẩm/đơn hàng       │   │',
        '  │  │   - Lịch sử đơn hàng           │   │    - Quản lý người dùng/coupon        │   │',
        '  │  │   - Vòng quay may mắn          │   │    - Live Chat & Liên hệ              │   │',
        '  │  └──────────────────────────────┘   └──────────────────────────────────────┘   │',
        '  └──────────────────────────────────────────────────────────────────────────────────┘',
        '                          │ RESTful API (HTTP/HTTPS)                                  ',
        '  ┌─────────────────────────────────────────────────────────────────────────────────┐',
        '  │                      TẦNG NGHIỆP VỤ (Business Logic Layer)                     │',
        '  │                         Laravel 10 (PHP 8.x) – Backend API                     │',
        '  │  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────────────┐ │',
        '  │  │   Auth   │  │ Product  │  │  Order   │  │  Coupon  │  │  Payment Gateway │ │',
        '  │  │ Sanctum  │  │ Category │  │  Cart    │  │  Lucky   │  │  VNPay / MoMo    │ │',
        '  │  └──────────┘  └──────────┘  └──────────┘  └──────────┘  └──────────────────┘ │',
        '  │  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌───────────────────────────────┐  │',
        '  │  │  Review  │  │LiveChat  │  │ Contact  │  │  AI Chatbot (Gemini API)       │  │',
        '  │  └──────────┘  └──────────┘  └──────────┘  └───────────────────────────────┘  │',
        '  └──────────────────────────────────────────────────────────────────────────────────┘',
        '                          │ Eloquent ORM (MySQL)                                      ',
        '  ┌─────────────────────────────────────────────────────────────────────────────────┐',
        '  │                        TẦNG DỮ LIỆU (Data Layer)                               │',
        '  │                              MySQL Database                                      │',
        '  │   users │ categories │ products │ orders │ order_items │ reviews │ cart_items   │',
        '  │   coupons │ lucky_wheel_spins │ chat_sessions │ chat_messages │ contacts         │',
        '  │   product_images │ payment_settings │ homepage_settings                         │',
        '  └──────────────────────────────────────────────────────────────────────────────────┘',
    ]
    add_ascii_box(doc, arch_lines, size=8)
    add_figure_caption(doc, 'Hình 3.1 – Kiến trúc tổng thể hệ thống bán laptop trực tuyến')

    # ─── 3.2 Sơ đồ tuần tự (Sequence Diagrams) ───
    add_heading(doc, '3.2. Thiết kế Sơ đồ tuần tự (Sequence Diagram)', level=2, size=14, color=(31, 56, 100))
    add_para(doc, 'Sơ đồ tuần tự mô tả luồng tương tác giữa các thành phần trong hệ thống theo trình tự thời gian. Dưới đây là các sơ đồ tuần tự cho những chức năng quan trọng nhất:')

    # SD 1: Đăng nhập
    add_heading(doc, '3.2.1. Sơ đồ tuần tự – Đăng nhập hệ thống', level=3, size=12, color=(70, 100, 160))
    sd1_lines = [
        '   User          Frontend         Backend API       Database',
        '    │                │                 │                │',
        '    │──Nhập email,──>│                 │                │',
        '    │  mật khẩu      │                 │                │',
        '    │                │──POST /login──>│                │',
        '    │                │                 │──SELECT users─>│',
        '    │                │                 │  WHERE email   │',
        '    │                │                 │<──User record──│',
        '    │                │                 │                │',
        '    │                │                 │─[Kiểm tra bcrypt hash]',
        '    │                │                 │                │',
        '    │                │<──200 OK + ─────│                │',
        '    │                │  Sanctum Token  │                │',
        '    │                │                 │                │',
        '    │<──Lưu token,───│                 │                │',
        '    │  redirect home │                 │                │',
        '    │                │                 │                │',
        '    │         [Lỗi: sai email/mật khẩu]                │',
        '    │                │<──401 Unauthorized──────────────│',
        '    │<──Hiển thị lỗi─│                 │                │',
    ]
    add_ascii_box(doc, sd1_lines, size=8.5)
    add_figure_caption(doc, 'Hình 3.2 – Sơ đồ tuần tự: Đăng nhập hệ thống')

    # SD 2: Đặt hàng
    add_heading(doc, '3.2.2. Sơ đồ tuần tự – Đặt hàng (Checkout)', level=3, size=12, color=(70, 100, 160))
    sd2_lines = [
        '   User          Frontend         Backend API       Database',
        '    │                │                 │                │',
        '    │──Xem giỏ hàng─>│                 │                │',
        '    │                │──GET /cart──── >│                │',
        '    │                │                 │──SELECT cart──>│',
        '    │                │<──Danh sách ─── │<──cart_items───│',
        '    │                │  giỏ hàng       │                │',
        '    │                │                 │                │',
        '    │──Nhập địa chỉ,─>│                │                │',
        '    │  chọn TT toán  │                 │                │',
        '    │──Nhấn "Đặt" ──>│                 │                │',
        '    │                │──POST /checkout>│                │',
        '    │                │  {address, pay} │                │',
        '    │                │                 │──SELECT stock──>│',
        '    │                │                 │<──Số lượng tồn─│',
        '    │                │                 │─[Kiểm tra tồn kho]',
        '    │                │                 │──BEGIN TRANS───>│',
        '    │                │                 │──INSERT orders─>│',
        '    │                │                 │──INSERT order_items>│',
        '    │                │                 │──UPDATE products│',
        '    │                │                 │  quantity-=qty  │',
        '    │                │                 │──DELETE cart────>│',
        '    │                │                 │──COMMIT─────── >│',
        '    │                │<──201 + order──>│                │',
        '    │<──Hiển thị────>│                 │                │',
        '    │  xác nhận đơn  │                 │                │',
        '    │                │                 │                │',
        '    │       [Hết hàng: ROLLBACK, báo lỗi]               │',
    ]
    add_ascii_box(doc, sd2_lines, size=8.5)
    add_figure_caption(doc, 'Hình 3.3 – Sơ đồ tuần tự: Đặt hàng (Checkout)')

    # SD 3: Đánh giá sản phẩm
    add_heading(doc, '3.2.3. Sơ đồ tuần tự – Đánh giá sản phẩm', level=3, size=12, color=(70, 100, 160))
    sd3_lines = [
        '   User          Frontend         Backend API       Database',
        '    │                │                 │                │',
        '    │──Xem lịch sử──>│                 │                │',
        '    │  đơn hàng      │──GET /orders──>│                │',
        '    │                │<──Danh sách ─── │                │',
        '    │                │  đơn hàng       │                │',
        '    │                │                 │                │',
        '    │──Chọn đơn ────>│                 │                │',
        '    │  đã nhận hàng  │                 │                │',
        '    │──Nhấn "Đánh"──>│                 │                │',
        '    │  giá"          │──Hiển thị form─>│                │',
        '    │<──Form đánh ───│                 │                │',
        '    │   giá (sao,    │                 │                │',
        '    │   nhận xét)    │                 │                │',
        '    │──Gửi đánh giá─>│                 │                │',
        '    │                │──POST /review──>│                │',
        '    │                │                 │──SELECT orders>│',
        '    │                │                 │ WHERE status=  │',
        '    │                │                 │ completed      │',
        '    │                │                 │<──Confirm──────│',
        '    │                │                 │──INSERT reviews>│',
        '    │                │                 │──UPDATE products│',
        '    │                │                 │  avg_rating    │',
        '    │                │<──200 OK───────>│                │',
        '    │<──Thông báo ───│                 │                │',
        '    │  thành công    │                 │                │',
        '    │                │                 │                │',
        '    │    [Đơn chưa giao: 403 Forbidden]│                │',
    ]
    add_ascii_box(doc, sd3_lines, size=8.5)
    add_figure_caption(doc, 'Hình 3.4 – Sơ đồ tuần tự: Đánh giá sản phẩm')

    # SD 4: Admin quản lý sản phẩm
    add_heading(doc, '3.2.4. Sơ đồ tuần tự – Admin quản lý sản phẩm', level=3, size=12, color=(70, 100, 160))
    sd4_lines = [
        '  Admin        Admin Panel      Backend API       Database',
        '    │                │                 │                │',
        '    │──Truy cập─────>│                 │                │',
        '    │  Quản lý SP    │──GET /products─>│                │',
        '    │                │                 │──SELECT all────>│',
        '    │                │<──Danh sách SP──│<──products──── │',
        '    │                │                 │                │',
        '    │   [THÊM SẢN PHẨM]                │                │',
        '    │──Nhập thông───>│                 │                │',
        '    │  tin SP + ảnh  │──POST /products>│                │',
        '    │                │  + images       │──Upload ảnh────>│',
        '    │                │                 │──INSERT products>│',
        '    │                │                 │──INSERT product_images',
        '    │                │<──201 Created───│                │',
        '    │<──Thông báo ───│                 │                │',
        '    │  thành công    │                 │                │',
        '    │                │                 │                │',
        '    │   [SỬA SẢN PHẨM]                 │                │',
        '    │──Chọn SP,─────>│                 │                │',
        '    │  cập nhật      │──PUT /products/{id}>│            │',
        '    │                │                 │──UPDATE products>│',
        '    │                │<──200 OK───────>│                │',
        '    │                │                 │                │',
        '    │   [XÓA SẢN PHẨM]                 │                │',
        '    │──Chọn SP,─────>│                 │                │',
        '    │  xác nhận xóa  │──DELETE /products/{id}>│         │',
        '    │                │                 │──DELETE products>│',
        '    │                │<──200 OK───────>│                │',
    ]
    add_ascii_box(doc, sd4_lines, size=8.5)
    add_figure_caption(doc, 'Hình 3.5 – Sơ đồ tuần tự: Admin quản lý sản phẩm')

    # SD 5: Vòng quay may mắn
    add_heading(doc, '3.2.5. Sơ đồ tuần tự – Vòng quay may mắn (Lucky Wheel)', level=3, size=12, color=(70, 100, 160))
    sd5_lines = [
        '   User          Frontend         Backend API       Database',
        '    │                │                 │                │',
        '    │──Truy cập─────>│                 │                │',
        '    │  Lucky Wheel   │                 │                │',
        '    │──Nhấn "Quay"──>│                 │                │',
        '    │                │──POST ─────────>│                │',
        '    │                │  /lucky-wheel   │──SELECT user──>│',
        '    │                │  /spin          │  last_spin     │',
        '    │                │                 │<──Last spin────│',
        '    │                │                 │─[Kiểm tra giới hạn quay]',
        '    │                │                 │                │',
        '    │                │                 │─[Random phần thưởng]',
        '    │                │                 │  (coupon/không trúng)',
        '    │                │                 │──INSERT lucky──>│',
        '    │                │                 │  _wheel_spins  │',
        '    │                │                 │  (nếu trúng)   │',
        '    │                │                 │──INSERT coupons>│',
        '    │                │<──Prize result──│                │',
        '    │<──Animation ───│                 │                │',
        '    │  + thông báo   │                 │                │',
        '    │  phần thưởng   │                 │                │',
        '    │                │                 │                │',
        '    │    [Đã quay hôm nay: 429 Too Many Requests]       │',
    ]
    add_ascii_box(doc, sd5_lines, size=8.5)
    add_figure_caption(doc, 'Hình 3.6 – Sơ đồ tuần tự: Vòng quay may mắn')

    # SD 6: Live Chat
    add_heading(doc, '3.2.6. Sơ đồ tuần tự – Live Chat hỗ trợ khách hàng', level=3, size=12, color=(70, 100, 160))
    sd6_lines = [
        '   User          Frontend         Backend API       Database    Admin Panel',
        '    │                │                 │                │             │',
        '    │──Mở Chat──────>│                 │                │             │',
        '    │                │──POST ─────────>│                │             │',
        '    │                │  /live-chat/    │──INSERT ──────>│             │',
        '    │                │  start          │  chat_sessions │             │',
        '    │                │<──session_id────│<──session──────│             │',
        '    │                │                 │                │             │',
        '    │──Nhập tin──────>│                │                │             │',
        '    │  nhắn          │──POST /live-chat│                │             │',
        '    │                │  /{id}/messages>│──INSERT ──────>│             │',
        '    │                │                 │  chat_messages │             │',
        '    │                │<──200 OK───────>│                │             │',
        '    │                │                 │                │             │',
        '    │                │  [Polling / WebSocket]           │             │',
        '    │                │                 │<──GET sessions─────────────>│',
        '    │                │                 │──messages─────>│             │',
        '    │                │                 │<──messages─────│             │',
        '    │                │                 │──────Hiển thị tin nhắn────>│',
        '    │                │                 │                │             │',
        '    │                │                 │<──Admin gửi────────────────│',
        '    │                │                 │  phản hồi      │             │',
        '    │                │                 │──INSERT ──────>│             │',
        '    │                │                 │  chat_messages │             │',
        '    │<──Hiển thị─────│<──Polling──────>│                │             │',
        '    │  phản hồi      │                 │                │             │',
    ]
    add_ascii_box(doc, sd6_lines, size=8)
    add_figure_caption(doc, 'Hình 3.7 – Sơ đồ tuần tự: Live Chat hỗ trợ khách hàng')

    # ─── 3.3 Sơ đồ quan hệ CSDL (ERD) ───
    add_heading(doc, '3.3. Sơ đồ Quan hệ Cơ sở dữ liệu (ERD & Data Dictionary)', level=2, size=14, color=(31, 56, 100))
    add_para(doc, 'Hệ thống sử dụng cơ sở dữ liệu MySQL với 15 bảng chính. Sơ đồ ERD thể hiện các quan hệ giữa các thực thể trong hệ thống:')

    # ERD Overview diagram
    erd_lines = [
        '                    ┌──────────────┐',
        '                    │    users      │',
        '                    │──────────────│',
        '                    │ id (PK)      │',
        '                    │ name         │',
        '                    │ email        │',
        '                    │ password     │',
        '                    │ phone        │',
        '                    │ address      │',
        '                    │ role         │',
        '                    └──────┬───────┘',
        '          ┌─────────────────┼──────────────────────────┐',
        '          │                 │                          │',
        '   ┌──────┴──────┐  ┌───────┴──────┐  ┌──────────────┴────┐',
        '   │  cart_items  │  │    orders    │  │     reviews       │',
        '   │─────────────│  │─────────────│  │───────────────────│',
        '   │ id (PK)     │  │ id (PK)     │  │ id (PK)           │',
        '   │ user_id(FK) │  │ user_id(FK) │  │ user_id(FK)       │',
        '   │ product_id  │  │ coupon_id   │  │ product_id(FK)    │',
        '   │ (FK)        │  │ (FK, null.) │  │ rating            │',
        '   │ quantity    │  │ total_price │  │ comment           │',
        '   └──────┬──────┘  │ payment_m.  │  │ created_at        │',
        '          │         │ status      │  └───────────────────┘',
        '          │         │ note        │          │',
        '          │         │ discount_am.│          │',
        '          │         └──────┬──────┘          │',
        '          │                │                  │',
        '          └────────────────┴──────────────────┘',
        '                           │',
        '                 ┌─────────┴──────────┐',
        '                 │    order_items      │',
        '                 │────────────────────│',
        '                 │ id (PK)            │',
        '                 │ order_id (FK)      │',
        '                 │ product_id (FK)    │',
        '                 │ quantity           │',
        '                 │ price              │',
        '                 └─────────┬──────────┘',
        '                           │',
        '                 ┌─────────┴──────────────────────────┐',
        '                 │           products                  │',
        '                 │────────────────────────────────────│',
        '                 │ id (PK)                            │',
        '                 │ name                               │',
        '                 │ description                        │',
        '                 │ price                              │',
        '                 │ discount                           │',
        '                 │ quantity                           │',
        '                 │ category_id (FK)                   │',
        '                 │ image                              │',
        '                 │ avg_rating                         │',
        '                 └────────────────────────────────────┘',
        '                           │                │',
        '              ┌────────────┘                └─────────────┐',
        '   ┌──────────┴──────┐                     ┌──────────────┴────┐',
        '   │   categories    │                     │  product_images   │',
        '   │─────────────────│                     │───────────────────│',
        '   │ id (PK)         │                     │ id (PK)           │',
        '   │ name            │                     │ product_id (FK)   │',
        '   └─────────────────┘                     │ image_url         │',
        '                                           └───────────────────┘',
    ]
    add_ascii_box(doc, erd_lines, size=7.5)
    add_figure_caption(doc, 'Hình 3.8 – Sơ đồ ERD (Phần 1): Quan hệ các bảng chính (users, products, orders)')

    # ERD part 2: coupon, chat, lucky wheel
    erd2_lines = [
        '   ┌──────────────────┐          ┌─────────────────────────────────────┐',
        '   │    coupons        │          │          lucky_wheel_spins           │',
        '   │──────────────────│          │─────────────────────────────────────│',
        '   │ id (PK)          │          │ id (PK)                             │',
        '   │ code (UNIQUE)    │          │ user_id (FK → users)                │',
        '   │ type (percent/   │          │ prize_name                          │',
        '   │       fixed)     │          │ coupon_code                         │',
        '   │ value            │          │ created_at                          │',
        '   │ min_order_value  │          │ updated_at                          │',
        '   │ max_discount     │          └─────────────────────────────────────┘',
        '   │ usage_limit      │',
        '   │ used_count       │          ┌─────────────────────────────────────┐',
        '   │ valid_from       │          │          chat_sessions               │',
        '   │ valid_until      │          │─────────────────────────────────────│',
        '   │ is_active        │          │ id (PK)                             │',
        '   └──────────────────┘          │ guest_id                            │',
        '          │ FK                   │ status (active/closed)              │',
        '          │                      │ created_at / updated_at             │',
        '   ┌──────┴──────────────┐       └──────────────┬──────────────────────┘',
        '   │      orders          │                      │ FK (1:N)',
        '   │ coupon_id (FK,null) │       ┌──────────────┴──────────────────────┐',
        '   │ discount_amount     │       │          chat_messages               │',
        '   └─────────────────────┘       │─────────────────────────────────────│',
        '                                 │ id (PK)                             │',
        '                                 │ chat_session_id (FK)                │',
        '                                 │ sender_type (user/admin)            │',
        '                                 │ message                             │',
        '                                 │ created_at / updated_at             │',
        '                                 └─────────────────────────────────────┘',
        '',
        '   ┌──────────────────────┐      ┌─────────────────────────────────────┐',
        '   │     contacts          │      │        payment_settings              │',
        '   │──────────────────────│      │─────────────────────────────────────│',
        '   │ id (PK)              │      │ id (PK)                             │',
        '   │ name                 │      │ setting_key (UNIQUE)                │',
        '   │ email                │      │ setting_value                       │',
        '   │ phone                │      │ created_at / updated_at             │',
        '   │ message              │      └─────────────────────────────────────┘',
        '   │ status               │',
        '   │ admin_note           │      ┌─────────────────────────────────────┐',
        '   │ created_at           │      │       homepage_settings              │',
        '   └──────────────────────┘      │─────────────────────────────────────│',
        '                                 │ id (PK)                             │',
        '                                 │ setting_key (UNIQUE)                │',
        '                                 │ setting_value (longText)            │',
        '                                 │ created_at / updated_at             │',
        '                                 └─────────────────────────────────────┘',
    ]
    add_ascii_box(doc, erd2_lines, size=7.5)
    add_figure_caption(doc, 'Hình 3.9 – Sơ đồ ERD (Phần 2): Coupons, Chat, Contacts, Settings')

    # ─── 3.3.2 Data Dictionary ───
    add_heading(doc, '3.3.1. Mô tả cơ sở dữ liệu (Data Dictionary)', level=3, size=13, color=(70, 100, 160))
    add_para(doc, 'Dưới đây là mô tả chi tiết cấu trúc từng bảng trong cơ sở dữ liệu:')

    tables_info = [
        {
            'name': 'Bảng users – Thông tin người dùng',
            'fig': 'Bảng 3.1',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3), Cm(3), Cm(7)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính tự tăng'),
                ('name', 'VARCHAR(100)', 'NOT NULL', 'Tên đầy đủ của người dùng'),
                ('email', 'VARCHAR(100)', 'NOT NULL, UNIQUE', 'Email đăng nhập (duy nhất)'),
                ('password', 'VARCHAR(255)', 'NOT NULL', 'Mật khẩu được mã hóa bcrypt'),
                ('phone', 'VARCHAR(20)', 'NULL', 'Số điện thoại liên hệ'),
                ('address', 'TEXT', 'NULL', 'Địa chỉ nhận hàng mặc định'),
                ('role', "ENUM('user','admin')", "DEFAULT 'user'", 'Vai trò trong hệ thống'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Thời điểm tạo tài khoản'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm cập nhật lần cuối'),
            ]
        },
        {
            'name': 'Bảng categories – Danh mục sản phẩm',
            'fig': 'Bảng 3.2',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3), Cm(3), Cm(7)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính tự tăng'),
                ('name', 'VARCHAR(100)', 'NOT NULL', 'Tên danh mục (VD: Gaming, Văn phòng)'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Thời điểm tạo'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm cập nhật'),
            ]
        },
        {
            'name': 'Bảng products – Thông tin sản phẩm laptop',
            'fig': 'Bảng 3.3',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính tự tăng'),
                ('name', 'VARCHAR(200)', 'NOT NULL', 'Tên sản phẩm laptop'),
                ('description', 'TEXT', 'NOT NULL', 'Mô tả chi tiết thông số kỹ thuật'),
                ('price', 'DECIMAL(15,2)', 'NOT NULL', 'Giá gốc của sản phẩm (VND)'),
                ('discount', 'INT', 'NULL', 'Phần trăm giảm giá (0–100)'),
                ('quantity', 'INT', 'NOT NULL', 'Số lượng tồn kho'),
                ('category_id', 'BIGINT UNSIGNED', 'FK → categories', 'ID danh mục sản phẩm'),
                ('image', 'VARCHAR(255)', 'NOT NULL', 'Đường dẫn ảnh đại diện chính'),
                ('avg_rating', 'DECIMAL(3,2)', 'NULL', 'Điểm đánh giá trung bình (1.00–5.00)'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Ngày thêm sản phẩm'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Ngày cập nhật lần cuối'),
            ]
        },
        {
            'name': 'Bảng product_images – Hình ảnh sản phẩm (đa ảnh)',
            'fig': 'Bảng 3.4',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3), Cm(3), Cm(7)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('product_id', 'BIGINT UNSIGNED', 'FK → products', 'Sản phẩm liên kết'),
                ('image_url', 'VARCHAR(500)', 'NOT NULL', 'URL ảnh sản phẩm bổ sung'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Thời điểm upload ảnh'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm cập nhật'),
            ]
        },
        {
            'name': 'Bảng orders – Đơn hàng',
            'fig': 'Bảng 3.5',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3.5), Cm(4), Cm(3), Cm(5.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('user_id', 'BIGINT UNSIGNED', 'FK → users', 'Người dùng đặt hàng'),
                ('total_price', 'DECIMAL(15,2)', 'NOT NULL', 'Tổng tiền đơn hàng (sau giảm giá)'),
                ('payment_method', "ENUM('COD','bank_transfer','vnpay','momo')", 'NOT NULL', 'Phương thức thanh toán'),
                ('status', "ENUM('pending','processing','shipping','completed','cancelled')", "DEFAULT 'pending'", 'Trạng thái đơn hàng'),
                ('coupon_id', 'BIGINT UNSIGNED', 'FK → coupons, NULL', 'Coupon đã áp dụng'),
                ('discount_amount', 'DECIMAL(15,2)', "DEFAULT 0", 'Số tiền được giảm'),
                ('note', 'TEXT', 'NULL', 'Ghi chú từ người dùng'),
                ('created_at', 'TIMESTAMP', 'NOT NULL', 'Ngày đặt hàng'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Ngày cập nhật trạng thái'),
            ]
        },
        {
            'name': 'Bảng order_items – Chi tiết đơn hàng',
            'fig': 'Bảng 3.6',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3.5), Cm(3.5), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('order_id', 'BIGINT UNSIGNED', 'FK → orders', 'Đơn hàng chứa sản phẩm này'),
                ('product_id', 'BIGINT UNSIGNED', 'FK → products', 'Sản phẩm được mua'),
                ('quantity', 'INT', 'NOT NULL', 'Số lượng đặt mua'),
                ('price', 'DECIMAL(15,2)', 'NOT NULL', 'Giá tại thời điểm đặt hàng'),
            ]
        },
        {
            'name': 'Bảng cart_items – Giỏ hàng',
            'fig': 'Bảng 3.7',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3), Cm(3.5), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('user_id', 'BIGINT UNSIGNED', 'FK → users', 'Người dùng sở hữu giỏ hàng'),
                ('product_id', 'BIGINT UNSIGNED', 'FK → products', 'Sản phẩm trong giỏ'),
                ('quantity', 'INT', 'NOT NULL', 'Số lượng sản phẩm trong giỏ'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Thời điểm thêm vào giỏ'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm cập nhật số lượng'),
            ]
        },
        {
            'name': 'Bảng reviews – Đánh giá sản phẩm',
            'fig': 'Bảng 3.8',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3), Cm(3), Cm(7)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('user_id', 'BIGINT UNSIGNED', 'FK → users', 'Người dùng đánh giá'),
                ('product_id', 'BIGINT UNSIGNED', 'FK → products', 'Sản phẩm được đánh giá'),
                ('rating', 'INT', 'NOT NULL, 1–5', 'Số sao đánh giá (1 đến 5)'),
                ('comment', 'TEXT', 'NOT NULL', 'Nội dung nhận xét chi tiết'),
                ('created_at', 'TIMESTAMP', 'NOT NULL', 'Ngày viết đánh giá'),
            ]
        },
        {
            'name': 'Bảng coupons – Mã giảm giá',
            'fig': 'Bảng 3.9',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3.5), Cm(3.5), Cm(3), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('code', 'VARCHAR(255)', 'NOT NULL, UNIQUE', 'Mã coupon duy nhất'),
                ('type', "ENUM('percent','fixed')", 'NOT NULL', 'Loại giảm giá: phần trăm hoặc cố định'),
                ('value', 'DECIMAL(15,2)', 'NOT NULL', 'Giá trị giảm giá'),
                ('min_order_value', 'DECIMAL(15,2)', "DEFAULT 0", 'Giá trị đơn hàng tối thiểu để áp dụng'),
                ('max_discount', 'DECIMAL(15,2)', 'NULL', 'Mức giảm tối đa (cho type=percent)'),
                ('usage_limit', 'INT', 'NULL', 'Số lần dùng tối đa (NULL = không giới hạn)'),
                ('used_count', 'INT', "DEFAULT 0", 'Số lần đã được sử dụng'),
                ('valid_from', 'TIMESTAMP', 'NULL', 'Thời điểm bắt đầu hiệu lực'),
                ('valid_until', 'TIMESTAMP', 'NULL', 'Thời điểm hết hạn'),
                ('is_active', 'TINYINT(1)', "DEFAULT 1", 'Trạng thái kích hoạt coupon'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Ngày tạo coupon'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Ngày cập nhật'),
            ]
        },
        {
            'name': 'Bảng lucky_wheel_spins – Lịch sử quay may mắn',
            'fig': 'Bảng 3.10',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3), Cm(3.5), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('user_id', 'BIGINT UNSIGNED', 'FK → users', 'Người dùng quay'),
                ('prize_name', 'VARCHAR(255)', 'NULL', 'Tên phần thưởng nhận được'),
                ('coupon_code', 'VARCHAR(255)', 'NULL', 'Mã coupon trúng thưởng (nếu có)'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Thời điểm quay'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm cập nhật'),
            ]
        },
        {
            'name': 'Bảng chat_sessions – Phiên hội thoại live chat',
            'fig': 'Bảng 3.11',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3), Cm(3.5), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('guest_id', 'VARCHAR(255)', 'NOT NULL', 'Session ID ẩn danh của khách'),
                ('status', 'VARCHAR(255)', "DEFAULT 'active'", 'Trạng thái: active / closed'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Thời điểm mở phiên chat'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm cập nhật'),
            ]
        },
        {
            'name': 'Bảng chat_messages – Tin nhắn live chat',
            'fig': 'Bảng 3.12',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3.5), Cm(3), Cm(3.5), Cm(6)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('chat_session_id', 'BIGINT UNSIGNED', 'FK → chat_sessions', 'Phiên hội thoại chứa tin nhắn'),
                ('sender_type', 'VARCHAR(255)', "NOT NULL ('user'/'admin')", 'Loại người gửi'),
                ('message', 'TEXT', 'NOT NULL', 'Nội dung tin nhắn'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Thời điểm gửi tin nhắn'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Thời điểm cập nhật'),
            ]
        },
        {
            'name': 'Bảng contacts – Yêu cầu liên hệ',
            'fig': 'Bảng 3.13',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3), Cm(3.5), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('name', 'VARCHAR(255)', 'NOT NULL', 'Tên người liên hệ'),
                ('email', 'VARCHAR(255)', 'NULL', 'Email của người liên hệ'),
                ('phone', 'VARCHAR(255)', 'NULL', 'Số điện thoại liên hệ'),
                ('message', 'TEXT', 'NOT NULL', 'Nội dung yêu cầu/hỏi đáp'),
                ('status', "ENUM('new','processing','resolved')", "DEFAULT 'new'", 'Trạng thái xử lý yêu cầu'),
                ('admin_note', 'TEXT', 'NULL', 'Ghi chú của admin khi xử lý'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Ngày gửi yêu cầu'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Ngày cập nhật'),
            ]
        },
        {
            'name': 'Bảng payment_settings – Cấu hình thanh toán',
            'fig': 'Bảng 3.14',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3), Cm(3.5), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('setting_key', 'VARCHAR(120)', 'NOT NULL, UNIQUE', 'Tên key cấu hình (vnpay_url, momo_key...)'),
                ('setting_value', 'TEXT', 'NULL', 'Giá trị cấu hình tương ứng'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Ngày tạo cấu hình'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Ngày cập nhật cấu hình'),
            ]
        },
        {
            'name': 'Bảng homepage_settings – Cài đặt giao diện trang chủ',
            'fig': 'Bảng 3.15',
            'cols': ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'],
            'widths': [Cm(3), Cm(3), Cm(3.5), Cm(6.5)],
            'rows': [
                ('id', 'BIGINT UNSIGNED', 'PK, AUTO', 'Khóa chính'),
                ('setting_key', 'VARCHAR(120)', 'NOT NULL, UNIQUE', 'Tên key cài đặt (banner_url, slider...)'),
                ('setting_value', 'LONGTEXT', 'NULL', 'Giá trị cài đặt (có thể lưu JSON dài)'),
                ('created_at', 'TIMESTAMP', 'NULL', 'Ngày tạo cài đặt'),
                ('updated_at', 'TIMESTAMP', 'NULL', 'Ngày cập nhật cài đặt'),
            ]
        },
    ]

    for tbl_info in tables_info:
        add_heading(doc, tbl_info['name'], level=4, size=12, color=(31, 56, 100), space_before=10, space_after=4)
        ncols = len(tbl_info['cols'])
        table = doc.add_table(rows=len(tbl_info['rows'])+1, cols=ncols)
        table.style = 'Table Grid'
        set_table_borders(table)
        add_table_header_row(table, tbl_info['cols'])
        for i, row_data in enumerate(tbl_info['rows']):
            row = table.rows[i+1]
            for j, cell_val in enumerate(row_data):
                row.cells[j].text = ''
                p = row.cells[j].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(cell_val)
                is_pk_fk = 'PK' in cell_val or 'FK' in cell_val or 'NOT NULL' in cell_val
                set_font(r, bold=(j == 0), size=10)
                if 'PK' in cell_val:
                    r.font.color.rgb = RGBColor(180, 20, 20)
                elif 'FK' in cell_val:
                    r.font.color.rgb = RGBColor(0, 80, 160)
            # Light shading for alternating rows
            if i % 2 == 1:
                for j in range(ncols):
                    tcPr = table.rows[i+1].cells[j]._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'F5F8FF')
                    tcPr.append(shd)
        for j, w in enumerate(tbl_info['widths']):
            for row in table.rows:
                row.cells[j].width = w
        add_figure_caption(doc, f"{tbl_info['fig']} – Cấu trúc bảng: {tbl_info['name'].split('–')[0].strip()}")
        doc.add_paragraph()

    # ─── 3.4 Cơ sở lý thuyết về Công nghệ cốt lõi ───
    add_heading(doc, '3.4. Cơ sở lý thuyết về Công nghệ cốt lõi', level=2, size=14, color=(31, 56, 100))

    tech_data = [
        ('Laravel 10 (PHP 8.x)', 'Framework PHP theo mô hình MVC. Cung cấp Eloquent ORM, Sanctum authentication, middleware, queue, RESTful routing. Xây dựng toàn bộ Backend API của hệ thống.'),
        ('Next.js 13+ (React)', 'Framework React hỗ trợ Server-Side Rendering (SSR) và Static Site Generation (SSG). Dùng cho Frontend người dùng và Admin Panel. Hỗ trợ routing động, API routes.'),
        ('MySQL 8.x', 'Hệ quản trị cơ sở dữ liệu quan hệ. Lưu trữ toàn bộ dữ liệu hệ thống với ràng buộc toàn vẹn (FK, UNIQUE, NOT NULL).'),
        ('Laravel Sanctum', 'Package xác thực API token-based cho SPA. Phát hành token sau đăng nhập, kiểm tra auth qua middleware cho mỗi request.'),
        ('React Context API', 'Quản lý trạng thái toàn cục phía frontend (user session, giỏ hàng, thông báo) không cần Redux.'),
        ('VNPay / MoMo API', 'Tích hợp cổng thanh toán điện tử phổ biến tại Việt Nam. Hệ thống redirect sang cổng thanh toán và xử lý callback.'),
        ('Gemini API (Google AI)', 'AI Chatbot tích hợp để tư vấn sản phẩm tự động. Backend gọi Gemini API và trả lời câu hỏi của khách hàng.'),
        ('RESTful API', 'Chuẩn thiết kế API với các phương thức GET/POST/PUT/DELETE, trả về JSON. Dễ dàng tích hợp với bất kỳ client nào.'),
    ]

    table = doc.add_table(rows=len(tech_data)+1, cols=2)
    table.style = 'Table Grid'
    set_table_borders(table)
    add_table_header_row(table, ['Công nghệ', 'Mô tả & Ứng dụng trong dự án'])
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(11.5)
    for i, (tech, desc) in enumerate(tech_data):
        row = table.rows[i+1]
        row.cells[0].text = ''
        row.cells[1].text = ''
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(tech)
        set_font(r0, bold=True, size=10)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(desc)
        set_font(r1, size=10)
        if i % 2 == 1:
            for j in range(2):
                tcPr = row.cells[j]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F5F8FF')
                tcPr.append(shd)
    add_figure_caption(doc, 'Bảng 3.16 – Công nghệ cốt lõi và ứng dụng trong hệ thống')

    # ─── 3.5 Cấu trúc thư mục dự án ───
    add_heading(doc, '3.5. Cấu trúc thư mục dự án', level=2, size=14, color=(31, 56, 100))
    add_para(doc, 'Dự án được tổ chức thành 3 phần chính: Backend (Laravel), Frontend người dùng (Next.js), và Admin Panel (Next.js):')

    struct_lines = [
        '  project-root/',
        '  ├── be/                           # Backend Laravel',
        '  │   ├── app/',
        '  │   │   ├── Http/Controllers/     # Controllers xử lý API',
        '  │   │   │   ├── AuthController.php        # Đăng nhập/đăng ký',
        '  │   │   │   ├── ProductController.php     # Sản phẩm, đánh giá',
        '  │   │   │   ├── CartController.php        # Giỏ hàng',
        '  │   │   │   ├── OrderController.php       # Đặt hàng, lịch sử',
        '  │   │   │   ├── CouponController.php      # Áp dụng coupon',
        '  │   │   │   ├── LuckyWheelController.php  # Vòng quay may mắn',
        '  │   │   │   ├── LiveChatController.php    # Live chat',
        '  │   │   │   ├── ChatController.php        # AI Chatbot (Gemini)',
        '  │   │   │   ├── ContactController.php     # Form liên hệ',
        '  │   │   │   ├── VnPayController.php       # Thanh toán VNPay',
        '  │   │   │   ├── MomoController.php        # Thanh toán MoMo',
        '  │   │   │   └── Admin/                   # Controllers Admin',
        '  │   │   │       ├── OrderController.php   # Quản lý đơn hàng',
        '  │   │   │       ├── UserController.php    # Quản lý người dùng',
        '  │   │   │       ├── CouponController.php  # Quản lý coupon',
        '  │   │   │       ├── DashboardController.php # Thống kê',
        '  │   │   │       └── ContactController.php # Xử lý liên hệ',
        '  │   │   └── Models/                      # Eloquent Models',
        '  │   ├── database/migrations/             # Cấu trúc CSDL',
        '  │   └── routes/api.php                   # Định nghĩa API routes',
        '  │',
        '  ├── fe1/                          # Frontend Next.js (Người dùng)',
        '  │   ├── pages/                    # Các trang của ứng dụng',
        '  │   │   ├── index.js             # Trang chủ',
        '  │   │   ├── products/            # Trang sản phẩm',
        '  │   │   ├── cart.js              # Giỏ hàng',
        '  │   │   ├── checkout.js          # Đặt hàng',
        '  │   │   ├── orders.js            # Lịch sử đơn',
        '  │   │   └── profile.js           # Thông tin cá nhân',
        '  │   ├── components/              # Components tái sử dụng',
        '  │   ├── context/                 # React Context (auth, cart)',
        '  │   └── services/               # Gọi API functions',
        '  │',
        '  └── admin/                        # Admin Panel Next.js',
        '      ├── pages/',
        '      │   ├── dashboard.js          # Thống kê tổng quan',
        '      │   ├── products/             # Quản lý sản phẩm',
        '      │   ├── orders/               # Quản lý đơn hàng',
        '      │   ├── users/                # Quản lý người dùng',
        '      │   ├── coupons/              # Quản lý coupon',
        '      │   ├── live-chat/            # Quản lý chat',
        '      │   └── contacts/             # Quản lý liên hệ',
        '      └── components/',
    ]
    add_ascii_box(doc, struct_lines, size=8)
    add_figure_caption(doc, 'Hình 3.10 – Cấu trúc thư mục dự án')

    # ─── 3.6 Giải pháp triển khai một số Module khó ───
    add_heading(doc, '3.6. Giải pháp triển khai một số Module khó', level=2, size=14, color=(31, 56, 100))

    # Module 1: Payment
    add_heading(doc, '3.6.1. Module Thanh toán VNPay và MoMo', level=3, size=13, color=(70, 100, 160))
    add_para(doc, 'Hệ thống tích hợp hai cổng thanh toán phổ biến tại Việt Nam: VNPay và MoMo. Luồng xử lý như sau:')
    payment_flow = [
        '  User chọn thanh toán VNPay/MoMo',
        '       │',
        '       ▼',
        '  Backend tạo URL thanh toán (signed với secret key)',
        '       │',
        '       ▼',
        '  Frontend redirect sang cổng thanh toán VNPay/MoMo',
        '       │',
        '  User hoàn tất thanh toán trên cổng',
        '       │',
        '       ▼',
        '  Cổng thanh toán gọi callback URL (GET /vnpay/return)',
        '  hoặc IPN URL (POST /momo/ipn)',
        '       │',
        '       ▼',
        '  Backend xác thực chữ ký (hash), cập nhật trạng thái đơn',
        '       │',
        '       ▼',
        '  Frontend hiển thị kết quả thanh toán thành công/thất bại',
    ]
    add_ascii_box(doc, payment_flow, size=9)
    add_figure_caption(doc, 'Hình 3.11 – Luồng xử lý thanh toán VNPay/MoMo')

    # Module 2: AI Chatbot
    add_heading(doc, '3.6.2. Module AI Chatbot (Gemini API)', level=3, size=13, color=(70, 100, 160))
    add_para(doc, 'Hệ thống tích hợp Google Gemini API để tự động tư vấn sản phẩm. Khi người dùng đặt câu hỏi qua chatbot, Backend gọi Gemini API với prompt kết hợp danh sách sản phẩm hiện có, AI trả lời dựa trên ngữ cảnh thực tế của cửa hàng:')
    chatbot_flow = [
        '  User nhập câu hỏi → POST /chat',
        '       │',
        '       ▼',
        '  Backend lấy danh sách sản phẩm từ DB',
        '  (tên, giá, tồn kho, category)',
        '       │',
        '       ▼',
        '  Tạo prompt: "Bạn là AI tư vấn laptop của ShopLaptop.',
        '  Danh sách sản phẩm: [products_json].',
        '  Câu hỏi người dùng: [question]"',
        '       │',
        '       ▼',
        '  Gọi Google Gemini API với prompt trên',
        '       │',
        '       ▼',
        '  Trả response về Frontend → Hiển thị cho User',
    ]
    add_ascii_box(doc, chatbot_flow, size=9)
    add_figure_caption(doc, 'Hình 3.12 – Luồng xử lý AI Chatbot tư vấn sản phẩm')

    # Module 3: Vòng quay may mắn
    add_heading(doc, '3.6.3. Module Vòng quay may mắn (Lucky Wheel)', level=3, size=13, color=(70, 100, 160))
    add_para(doc, 'Module gamification giúp tăng tương tác người dùng. Người dùng đăng nhập có thể quay vòng quay 1 lần/ngày để nhận coupon giảm giá. Hệ thống kiểm tra lịch sử quay trong bảng lucky_wheel_spins và tự động tạo coupon nếu trúng thưởng.')
    lw_lines = [
        '  Cơ chế phân bổ phần thưởng:',
        '  ┌─────────────────────────────────────────────────────────────┐',
        '  │  Phần thưởng          │  Xác suất  │  Mô tả                │',
        '  │───────────────────────────────────────────────────────────── │',
        '  │  Coupon 5% (code tự động) │  20%    │  Giảm 5% tối đa 100k │',
        '  │  Coupon 10%             │  10%      │  Giảm 10% tối đa 200k│',
        '  │  Coupon 15%             │  5%       │  Giảm 15% tối đa 300k│',
        '  │  Coupon Fixed 50k       │  10%      │  Giảm cố định 50.000đ│',
        '  │  Không trúng           │  55%       │  Chúc may mắn lần sau│',
        '  └─────────────────────────────────────────────────────────────┘',
        '',
        '  Kiểm tra giới hạn: SELECT * FROM lucky_wheel_spins',
        '                     WHERE user_id = ? AND DATE(created_at) = TODAY',
        '  → Nếu đã quay: trả 429 Too Many Requests',
        '  → Nếu chưa quay: xử lý random → INSERT lucky_wheel_spins',
        '                   → INSERT coupons (nếu trúng) → trả kết quả',
    ]
    add_ascii_box(doc, lw_lines, size=8.5)
    add_figure_caption(doc, 'Hình 3.13 – Cơ chế phân bổ phần thưởng Vòng quay may mắn')

    # ─── 3.7 Mô hình luồng dữ liệu tổng thể ───
    add_heading(doc, '3.7. Mô hình luồng dữ liệu (Data Flow Diagram – DFD) tổng thể', level=2, size=14, color=(31, 56, 100))
    add_para(doc, 'Sơ đồ luồng dữ liệu mức 0 (Context Diagram) và mức 1 của hệ thống:')

    dfd_lines = [
        '  ╔══════════════════════════════════════════════════════════════════════╗',
        '  ║                  DFD MỨC 0 – CONTEXT DIAGRAM                       ║',
        '  ╠══════════════════════════════════════════════════════════════════════╣',
        '  ║                                                                      ║',
        '  ║   [Người dùng] ─────────────────────────────────────────────────    ║',
        '  ║      │  Nhập thông tin, đặt hàng,    ↑ Kết quả xử lý, thông báo    ║',
        '  ║      │  đánh giá, chat              │                               ║',
        '  ║      ▼                              │                               ║',
        '  ║   ┌──────────────────────────────────────────────────────────────┐  ║',
        '  ║   │           HỆ THỐNG BÁN LAPTOP TRỰC TUYẾN                   │  ║',
        '  ║   │    (Backend Laravel + Frontend Next.js + Admin Panel)        │  ║',
        '  ║   └──────────────────────────────────────────────────────────────┘  ║',
        '  ║      │  Thông tin SP, giá,           ↑ Dữ liệu lưu trữ             ║',
        '  ║      │  đơn hàng,...                │                               ║',
        '  ║      ▼                              │                               ║',
        '  ║   [Cơ sở dữ liệu MySQL]             │                               ║',
        '  ║      │                              │                               ║',
        '  ║   [Admin] ──→ [Quản lý & thống kê] ─┘                               ║',
        '  ║                                                                      ║',
        '  ║   [VNPay/MoMo] ←──────── Thanh toán ─────────────────────→ [User] ║',
        '  ║   [Gemini AI]  ←──────── Câu hỏi   ──────────────→ [Chatbot API] ║',
        '  ╚══════════════════════════════════════════════════════════════════════╝',
    ]
    add_ascii_box(doc, dfd_lines, size=8)
    add_figure_caption(doc, 'Hình 3.14 – Sơ đồ luồng dữ liệu mức 0 (Context Diagram)')


# ═══════════════════════════════════════════════════════════════
#  MAIN – Tạo Document
# ═══════════════════════════════════════════════════════════════

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)   # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    build_chapter2(doc)
    build_chapter3(doc)

    out_path = '/Users/nguyennghia/Desktop/workspace/ManhLaptop/Chuong2_Chuong3_BaoCao.docx'
    doc.save(out_path)
    print(f'✅ File đã được tạo: {out_path}')
    return out_path


if __name__ == '__main__':
    main()
